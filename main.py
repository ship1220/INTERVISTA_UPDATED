# main.py — Refactored async FastAPI application

import asyncio
import io
import json
import logging
import os
import re
import time
from datetime import datetime, timezone, timedelta
from pathlib import Path

from utils.logging_config import configure_app_logging, log_startup_banner

configure_app_logging()
from speech.transcription import transcribe_audio, _convert_audio
from fastapi import FastAPI, Request, Form, Depends, HTTPException, UploadFile, File, Body
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlalchemy import inspect, text
from sqlalchemy.orm import Session
from passlib.context import CryptContext
try:
    from sse_starlette.sse import EventSourceResponse
except Exception:
    EventSourceResponse = None
    # Defer import failure until SSE endpoints are used; log a warning at runtime
from database import Base, engine, SessionLocal
from models import (
    User,
    InterviewAttempt,
    SkillProgress,
    UserProfile,
    Interview,
    UserSkillProfileRow,
    Course,
    Module,
    ModuleAttempt,
    CourseResource,
)
from user_skill_profile import (
    BasicUserInfo,
    ResumeData,
    UserSkillProfile,
    UserSkillVector,
    InterviewRecord,
    ScoreBreakdown,
    create_user_profile,
    detect_weaknesses,
    update_skill_score,
    update_skill_vector,
    calculate_overall_score,
    record_interview_result,
    score_answer_structure,
    blend_dimension_score,
    aggregate_dimension_scores,
    build_dashboard_analytics,
)
from speech.transcription import normalize_transcript
from core.llm.llm_service import LLMService
from core.prompts.prompt_manager import PromptManager
from services.rag.retriever import get_retriever
from services.rag.rag_pipeline import get_or_create_rag_pipeline, initialize_rag
from core.chains.base_chain import InterviewQuestionChain, EvaluationChain, SummaryChain
from speech.transcription import (
    transcribe_audio,
    analyze_speech_delivery,
    compute_confidence_score,
    compute_overall_score,
    compute_recruiter_verdict,
)
from utils.rl_helpers import (
    get_state_id,
    calculate_reward,
)
from utils.bandit_logger import (
    log_bandit_state,
    log_reward_calculation,
    log_course_generation_decision,
    log_bandit_complete,
)
from utils.jd_analysis import analyze_resume_vs_jd
from services.rl.rl_service import ContextualBandit, INTERVIEW_ACTIONS, COURSE_ACTIONS
from models import QTable, UserState
from services.learning_resources.pipeline import ResourceRetrievalPipeline
from services.learning_resources.ranker import SemanticRanker

Base.metadata.create_all(bind=engine)


def ensure_database_schema(engine):
    inspector = inspect(engine)
    table_names = inspector.get_table_names()

    with engine.begin() as conn:
        if "courses" in table_names:
            course_cols = {col["name"] for col in inspector.get_columns("courses")}
            if "role" not in course_cols:
                conn.execute(text('ALTER TABLE courses ADD COLUMN role VARCHAR;'))
            if "created_at" not in course_cols:
                conn.execute(text('ALTER TABLE courses ADD COLUMN created_at TIMESTAMP;'))
            if "updated_at" not in course_cols:
                conn.execute(text('ALTER TABLE courses ADD COLUMN updated_at TIMESTAMP;'))

        if "modules" in table_names:
            module_cols = {col["name"] for col in inspector.get_columns("modules")}
            if "is_final" not in module_cols:
                conn.execute(text('ALTER TABLE modules ADD COLUMN is_final BOOLEAN DEFAULT FALSE;'))

        if "user_profiles" in table_names:
            profile_cols = {col["name"] for col in inspector.get_columns("user_profiles")}
            if "company_name" not in profile_cols:
                conn.execute(text('ALTER TABLE user_profiles ADD COLUMN company_name VARCHAR;'))
            if "job_description" not in profile_cols:
                conn.execute(text('ALTER TABLE user_profiles ADD COLUMN job_description TEXT;'))

        # Phase 2: Ensure course_resources table exists
        if "course_resources" not in table_names:
            from models import CourseResource as _CR
            _CR.__table__.create(bind=engine, checkfirst=True)
        else:
            cr_cols = {col["name"] for col in inspector.get_columns("course_resources")}
            if "explanation" not in cr_cols:
                conn.execute(text("ALTER TABLE course_resources ADD COLUMN explanation TEXT;"))


ensure_database_schema(engine)

import shutil

UPLOAD_DIR = "uploads"

# New architecture instances
llm_service = LLMService()
prompt_manager = PromptManager()
retriever = get_retriever()

# RAG Pipeline - initialized on startup
rag_pipeline = None

# Learning Resources Pipeline (Phase 2) - initialized on startup
resource_pipeline = None

question_chain = InterviewQuestionChain(llm_service, prompt_manager, retriever)
evaluation_chain = EvaluationChain(llm_service, prompt_manager, retriever)
summary_chain = SummaryChain(llm_service, prompt_manager, retriever)

def categorize_weak_topics(weak_topics: list[str]) -> dict[str, list[str]]:

    technical = []
    communication = []

    for topic in weak_topics:
        t = topic.lower()

        if any(k in t for k in [
            "sql","database","python","algorithm","data structure",
            "machine learning","statistics","api","system design"
        ]):
            technical.append(topic)

        elif any(k in t for k in [
            "explain","clarity","structure","communication",
            "example","confidence","detail","depth"
        ]):
            communication.append(topic)

        else:
            # fallback
            technical.append(topic)

    return {
        "Technical Skills": list(set(technical)),
        "Communication & Answer Quality": list(set(communication))
    }


def _save_interview_profile_fields(
    db: Session,
    user: User,
    *,
    role: str | None = None,
    level: str | None = None,
    company_name: str | None = None,
    job_description: str | None = None,
) -> UserProfile:
    """Persist optional interview targeting fields on UserProfile."""
    profile = get_or_create_user_profile(db, user)
    if role is not None and role.strip():
        profile.role_applied_for = role.strip()
    if level is not None and level.strip():
        profile.current_designation = level.strip()
    if company_name is not None:
        profile.company_name = company_name.strip() or None
    if job_description is not None:
        profile.job_description = job_description.strip() or None
    profile.updated_at = datetime.utcnow()
    db.add(profile)
    db.commit()
    return profile


async def _ensure_session_jd_analysis(session: dict, resume_text: str) -> dict:
    """Run resume vs JD analysis once per interview session."""
    if "jd_analysis" in session:
        return session["jd_analysis"]

    job_description = (session.get("job_description") or "").strip()
    if not job_description:
        analysis = {"matched_skills": [], "missing_skills": [], "ats_score": None}
    else:
        analysis = await analyze_resume_vs_jd(
            resume_text,
            job_description,
            llm_service=llm_service,
            prompt_manager=prompt_manager,
        )

    session["jd_analysis"] = analysis
    return analysis


def _interview_jd_payload(session: dict) -> dict:
    """Build JD-related fields for question/evaluation chains."""
    job_description = (session.get("job_description") or "").strip()
    analysis = session.get("jd_analysis") or {}
    if not job_description:
        return {
            "company_name": "",
            "job_description": "",
            "matched_skills": [],
            "missing_skills": [],
        }
    return {
        "company_name": (session.get("company_name") or "").strip(),
        "job_description": job_description,
        "matched_skills": analysis.get("matched_skills", []),
        "missing_skills": analysis.get("missing_skills", []),
    }


PLACEMENT_ROUND_ORDER = [
    "online_assessment",
    "technical",
    "group_discussion",
    "hr",
]

EXECUTABLE_ROUNDS = {"technical"}


def _normalize_interview_mode(mode: str | None) -> str:
    normalized = (mode or "placement_simulation").strip().lower()
    if normalized in ("individual_practice", "placement_simulation"):
        return normalized
    return "placement_simulation"


def _normalize_round(round_id: str | None) -> str:
    normalized = (round_id or "technical").strip().lower()
    if normalized in PLACEMENT_ROUND_ORDER:
        return normalized
    return "technical"


def _init_interview_session_fields(
    interview_mode: str = "placement_simulation",
    selected_round: str | None = None,
) -> dict:
    """Build interview_mode / current_round / round_order / round_status for a new session."""
    mode = _normalize_interview_mode(interview_mode)
    if mode == "individual_practice":
        round_id = _normalize_round(selected_round)
        return {
            "interview_mode": "individual_practice",
            "current_round": round_id,
            "round_order": [round_id],
            "round_status": {round_id: "active"},
        }
    return {
        "interview_mode": "placement_simulation",
        "current_round": "online_assessment",
        "round_order": list(PLACEMENT_ROUND_ORDER),
        "round_status": {
            "online_assessment": "active",
            "technical": "pending",
            "group_discussion": "pending",
            "hr": "pending",
        },
    }


def _advance_round(session: dict) -> str | None:
    """Mark the current round complete and activate the next one."""
    current = session.get("current_round")
    order = session.get("round_order", list(PLACEMENT_ROUND_ORDER))
    status = session.setdefault("round_status", {})
    if current:
        status[current] = "completed"
    try:
        idx = order.index(current)
    except ValueError:
        session["current_round"] = None
        return None
    if idx + 1 >= len(order):
        session["current_round"] = None
        return None
    nxt = order[idx + 1]
    session["current_round"] = nxt
    status[nxt] = "active"
    return nxt


def _skip_to_executable_round(session: dict) -> None:
    """Advance through unimplemented rounds until an executable one (entry routing only)."""
    while (
        session.get("current_round")
        and session["current_round"] not in EXECUTABLE_ROUNDS
    ):
        if _advance_round(session) is None:
            break


def _ensure_round_fields(session: dict) -> None:
    """Backfill round fields on sessions created before multi-round support."""
    if "current_round" in session:
        return
    if "current_stage" in session:
        session["current_round"] = session.pop("current_stage")
        session["round_order"] = session.pop("stage_order", list(PLACEMENT_ROUND_ORDER))
        session["round_status"] = session.pop("stage_status", {})
        session.setdefault("interview_mode", "placement_simulation")
        return
    session.update(_init_interview_session_fields("placement_simulation"))


def _prepare_interview_session(
    session: dict,
    *,
    interview_mode: str,
    selected_round: str | None = None,
) -> None:
    """Apply mode-specific entry routing so the session lands on the correct round."""
    mode = _normalize_interview_mode(interview_mode)
    session["interview_mode"] = mode
    if mode == "placement_simulation":
        _skip_to_executable_round(session)
    else:
        round_id = _normalize_round(selected_round or session.get("current_round"))
        session["current_round"] = round_id
        session["round_order"] = [round_id]
        session["round_status"] = {round_id: "active"}


# ===========================================================================
# APP INIT
# ===========================================================================
app = FastAPI()
BASE_DIR = Path(__file__).parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

logger = logging.getLogger(__name__)


@app.middleware("http")
async def log_http_requests(request: Request, call_next):
    """Minimal request log (replaces verbose uvicorn access lines)."""
    start = time.perf_counter()
    response = await call_next(request)
    elapsed_ms = (time.perf_counter() - start) * 1000
    path = request.url.path
    if path.startswith("/static"):
        return response
    print(
        f"  {request.method} {path} → {response.status_code} ({elapsed_ms:.0f}ms)",
        flush=True,
    )
    return response

# Use PBKDF2-SHA512 as primary (no 72-byte limit), keep bcrypt for backward compatibility
pwd_context = CryptContext(
    schemes=["pbkdf2_sha512", "bcrypt"],
    deprecated="auto"
)

# In-memory interview session store (swap for Redis in production)
interview_sessions: dict[str, dict] = {}

# In-memory resume text store (swap for Redis/DB in production)
resume_store: dict[str, str] = {}

# In-memory report store (latest report per user)
report_store: dict[str, dict] = {}


# ===========================================================================
# RAG PIPELINE STARTUP EVENT
# ===========================================================================
@app.on_event("startup")
async def startup_rag_pipeline():
    """Initialize RAG pipeline and learning resources pipeline."""

    global rag_pipeline, resource_pipeline
    from config.settings import settings

    log_startup_banner(settings.api_host, settings.api_port)

    try:
        rag_pipeline = get_or_create_rag_pipeline()
        asyncio.create_task(initialize_rag_async())
    except Exception:
        logger.exception("RAG pipeline startup failed")
        rag_pipeline = None

    # Initialize learning resources pipeline (Phase 2)
    try:
        vector_store = rag_pipeline.vector_store if rag_pipeline else None
        embedder = vector_store.embedder if vector_store else None
        ranker = SemanticRanker(embedder=embedder) if embedder else None

        resource_pipeline = ResourceRetrievalPipeline(
            llm_service=llm_service,
            prompt_manager=prompt_manager,
            youtube_api_key=settings.youtube_api_key or None,
            ranker=ranker,
            vector_store=vector_store,
        )
        logger.info("Learning resources pipeline initialized")
    except Exception:
        logger.exception("Learning resources pipeline startup failed")
        resource_pipeline = None


async def initialize_rag_async():
    """Initialize RAG asynchronously (errors only)."""
    if rag_pipeline is None:
        return

    try:
        await initialize_rag()
    except Exception:
        logger.exception("RAG pipeline async initialization failed")


# ===========================================================================
# AI SERVICE HELPERS
# ===========================================================================

def extract_json(text: str):
    import json
    import re

    if not text:
        raise ValueError("Empty LLM response")

    t = text.strip()

    # Remove common fenced-code markers (```json, ```), inline backtick wrappers, and language hints
    # Do this liberally to maximize chance of extracting the JSON block.
    t = re.sub(r'```\s*json', '', t, flags=re.IGNORECASE)
    t = re.sub(r'```', '', t)
    t = re.sub(r'`json\b', '', t, flags=re.IGNORECASE)
    # Remove remaining single backticks which often surround inline JSON
    t = t.replace('`', '')

    # Find the first balanced JSON object using brace counting
    start = t.find('{')
    if start == -1:
        raise ValueError("No JSON object found in response")

    depth = 0
    end = -1
    for idx in range(start, len(t)):
        ch = t[idx]
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                end = idx
                break

    if end == -1:
        raise ValueError("No matching closing brace for JSON object")

    json_str = t[start:end + 1]

    # Remove invalid control characters that break json.loads
    json_str = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)

    # Fix common LLM-caused trailing commas inside objects/arrays
    json_str = re.sub(r',\s*}', '}', json_str)
    json_str = re.sub(r',\s*]', ']', json_str)

    return json.loads(json_str, strict=False)


def safe_json_loads(json_str: str):
    """Safely parse JSON string by removing invalid control characters."""
    import json
    import re
    if not json_str:
        return None
    # Remove invalid control characters
    cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', json_str)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None


def _get_recent_interview_metrics(user_id: int, db: Session, lookback: int = 5) -> dict:
    """Return the most recent recommended action and persistence metrics for a user."""
    interviews = (
        db.query(Interview)
        .filter(Interview.user_id == user_id)
        .order_by(Interview.date.desc())
        .limit(lookback)
        .all()
    )

    last_action = None
    action_streak = 0
    previous_weak_topics = []
    previous_confidence = 0.0
    weak_history: list[tuple[str, ...]] = []

    for interview in interviews:
        try:
            report = json.loads(interview.report_json)
        except Exception:
            continue

        action = report.get("recommended_action")
        if not action:
            continue

        if last_action is None:
            last_action = action
            action_streak = 1
        elif action == last_action:
            action_streak += 1
        else:
            break

        weak_topics = report.get("weak_topics") or report.get("rl_metrics", {}).get("weak_topics", [])
        if isinstance(weak_topics, str):
            weak_topics = [weak_topics]
        if not previous_weak_topics and isinstance(weak_topics, list):
            previous_weak_topics = [str(t).strip() for t in weak_topics if str(t).strip()]

        if isinstance(weak_topics, list):
            normalized = tuple(sorted({str(t).strip().lower() for t in weak_topics if str(t).strip()}))
            if normalized:
                weak_history.append(normalized)

        if previous_confidence == 0.0:
            previous_confidence = report.get("voice_analysis", {}).get("confidence_score", 0.0) or 0.0

    persistent_weak_topics = len(weak_history) >= 2 and weak_history[0] == weak_history[1]

    return {
        "last_action": last_action,
        "action_streak": action_streak,
        "previous_weak_topics": previous_weak_topics,
        "previous_confidence": previous_confidence,
        "persistent_weak_topics": persistent_weak_topics,
    }


def _prioritize_top_weak_topics(weak_topics: list[str], count: int = 3) -> list[str]:
    return [str(t).strip() for t in weak_topics if str(t).strip()][:count]



def normalize_questions_output(raw: str, count: int = 5) -> list[str]:
    questions = []

    if raw:
        try:
            parsed = extract_json(raw)
            if isinstance(parsed, list):
                questions = parsed
            elif isinstance(parsed, dict):
                if "questions" in parsed:
                    questions = parsed["questions"]
                elif "question" in parsed:
                    questions = parsed["question"]
                else:
                    for value in parsed.values():
                        if isinstance(value, list):
                            questions = value
                            break
        except Exception:
            # Fall back to regex and heuristic extraction
            matches = re.findall(r'"question"\s*:\s*"([^"]+)"', raw)
            if matches:
                questions = matches
            else:
                questions = [
                    line.strip().lstrip("0123456789.)- ")
                    for line in raw.splitlines()
                    if line.strip().endswith("?")
                ]

    cleaned_questions = []
    for q in questions:
        if isinstance(q, str):
            cleaned_questions.append(q.strip())
        elif isinstance(q, dict):
            if "question" in q:
                cleaned_questions.append(str(q["question"]).strip())
            elif "text" in q:
                cleaned_questions.append(str(q["text"]).strip())

    return [q for q in cleaned_questions if q][:count]


async def generate_interview_questions(role: str, level: str, count: int = 5, resume_text: str = "") -> list[str]:
    result = await question_chain.invoke(
        {
            "role": role,
            "level": level,
            "count": count,
            "resume_text": resume_text,
        }
    )

    if result.status != "success":
        return []

    try:
        return normalize_questions_output(result.output, count)
    except Exception:
        return []


async def evaluate_content(
    role: str,
    level: str,
    questions_answers: list,
    company_name: str = "",
    job_description: str = "",
) -> dict:
    answers = []
    weak_topics = []

    per_answer_dimensions: list[dict] = []

    for qa in questions_answers:
        question = qa.get("question", "")
        answer = normalize_transcript(qa.get("answer", ""))
        heuristics = score_answer_structure(answer, question)

        result = await evaluation_chain.invoke(
            {
                "role": role,
                "level": level,
                "question": question,
                "answer": answer,
                "company_name": company_name or "N/A",
                "job_description": job_description or "N/A",
            }
        )

        parsed = {}
        try:
            parsed = safe_json_loads(result.output) or {}
            if not isinstance(parsed, dict):
                parsed = {}
        except Exception:
            parsed = {}

        # Validate and extract score
        score = parsed.get("score", 0)
        try:
            score = float(score) if score else 0
        except (TypeError, ValueError):
            score = 0
        score = max(0, min(100, score))

        # Validate and extract lists
        answer_normalized = str(answer or "").strip().lower()
        if answer_normalized in ["", "(skipped)", "(no response)"]:
            strengths = ["No answer provided."]
        else:
            strengths = parsed.get("strengths", ["Answer attempted."])
            if not isinstance(strengths, list):
                strengths = ["Answer attempted."]
            strengths = [str(s).strip() for s in strengths if s][:3]
            if not strengths:
                strengths = ["Answer attempted."]

        weaknesses = parsed.get("weaknesses", ["Needs improvement."])
        if not isinstance(weaknesses, list):
            weaknesses = ["Needs improvement."]
        weaknesses = [str(w).strip() for w in weaknesses if w][:3]
        if not weaknesses:
            weaknesses = ["Needs improvement."]

        # Validate ideal_answer
        ideal_answer = parsed.get("ideal_answer", "Ideal answer unavailable.")
        if not isinstance(ideal_answer, str):
            ideal_answer = str(ideal_answer) if ideal_answer else "Ideal answer unavailable."

        # Validate weak_topics - MUST be list
        weak_topics_raw = parsed.get("weak_topics", [])
        if isinstance(weak_topics_raw, str):
            weak_topics_raw = [weak_topics_raw]
        if not isinstance(weak_topics_raw, list):
            weak_topics_raw = []
        weak_topics_list = [str(t).strip().lower() for t in weak_topics_raw if t]
        weak_topics_list = list(set(weak_topics_list))[:5]

        dimension_scores = {
            "relevance": blend_dimension_score(
                parsed.get("relevance_score"), heuristics["relevance"]
            ),
            "explanation_depth": blend_dimension_score(
                parsed.get("explanation_depth_score"), heuristics["explanation_depth"]
            ),
            "star_method": blend_dimension_score(
                parsed.get("star_method_score"), heuristics["star_method"]
            ),
            "structured_thinking": blend_dimension_score(
                parsed.get("structured_thinking_score"), heuristics["structured_thinking"]
            ),
            "problem_solving": blend_dimension_score(
                parsed.get("problem_solving_score"), heuristics["problem_solving"]
            ),
        }
        per_answer_dimensions.append(dimension_scores)
        
        # ============================================================
        # RL: EXTRACT CKFS METRICS (NEW)
        # ============================================================
        rl_metrics = {
            "C": max(0.0, min(1.0, float(parsed.get("C", 0.0)))),
            "K": max(0.0, min(1.0, float(parsed.get("K", 0.0)))),
            "F": max(0.0, min(1.0, float(parsed.get("F", 0.0)))),
            "S": max(0.0, min(1.0, float(parsed.get("S", 0.0))))
        }
        
        item = {
            "score": score,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "ideal_answer": ideal_answer,
            "weak_topics": weak_topics_list,
            "dimension_scores": dimension_scores,
            "_rl_metrics": rl_metrics,  # Internal use only, not sent to frontend
        }
        
        weak_topics.extend(weak_topics_list)
        answers.append(item)

    # Deduplicate and clean weak_topics
    weak_topics_final = list(set([t.strip().lower() for t in weak_topics if t]))
    weak_topics_final = [t for t in weak_topics_final if t][:10]

    scores = [a["score"] for a in answers] if answers else [0]
    avg_score = sum(scores) / len(scores) if scores else 0
    dim_agg = aggregate_dimension_scores(per_answer_dimensions)

    return {
        "answers": answers,
        "weak_topics": weak_topics_final,
        "overall_feedback": "",
        "aggregate": {
            "technical_score": round(avg_score),
            "communication_score": round(avg_score * 0.9),
            "overall_score": round(avg_score),
            "relevance_score": dim_agg.get("relevance", round(avg_score)),
            "depth_score": dim_agg.get("explanation_depth", round(avg_score * 0.9)),
            "star_method_score": dim_agg.get("star_method", round(avg_score * 0.85)),
            "structured_thinking_score": dim_agg.get(
                "structured_thinking", round(avg_score * 0.9)
            ),
            "problem_solving_score": dim_agg.get(
                "problem_solving", round(avg_score * 0.88)
            ),
        },
    }


async def evaluate_answers(role: str, questions_answers: list) -> dict:
    per_question = []

    for qa in questions_answers:
        question = qa.get("question", "")
        answer = qa.get("answer", "")

        result = await evaluation_chain.invoke(
            {
                "role": role,
                "level": "Junior",
                "question": question,
                "answer": answer,
            }
        )

        parsed = {}
        try:
            parsed = safe_json_loads(result.output) or {}
        except Exception:
            pass

        per_question.append({
            "question": question,
            "candidate_answer": answer,
            "feedback": parsed.get("strengths", []),
            "improved_answer": parsed.get("ideal_answer", ""),
        })

    return {
        "feedback_per_question": per_question,
        "improvement_tips": [
            "Practice structuring your answers with clear examples.",
            "Use the STAR method for behavioral responses.",
            "Focus on clarity, confidence, and relevance to the role.",
        ],
        "learning_resources": [
            {"topic": "Interview Preparation", "resource": "Practice common interview questions and structure answers clearly."}
        ],
    }


async def generate_performance_summary(report_data: dict) -> str:
    result = await summary_chain.invoke(
        {
            "role": report_data.get("candidate_profile", {}).get("role", "Candidate"),
            "score": report_data.get("overall_score", 0),
            "weak_topics": report_data.get("weak_topics", []),
            "attempted": len([
                a for a in report_data.get("detailed_answers", [])
                if a.get("transcript") not in ["", "(skipped)", "(no response)"]
            ]),
        }
    )

    return result.output or "Summary unavailable."


async def generate_resume_skill_profile(resume_text: str, role: str, level: str) -> dict:
    prompt = prompt_manager.get_prompt(
        "resume_skill_profile",
        resume_text=resume_text,
        role=role,
        experience_level=level,
    )

    raw = await llm_service.invoke(prompt, use_cache=False, json_mode=True)
    try:
        return extract_json(raw)
    except Exception:
        return {
            "identified_skills": [],
            "skill_gaps": [],
            "missing_for_role": [],
            "strength_percentage": 0,
            "improvement_suggestions": "Unable to parse resume insights.",
        }


# ===========================================================================
# DATABASE DEPENDENCY
# ===========================================================================
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ===========================================================================
# SPEECH ANALYSIS HELPERS
# ===========================================================================
def analyze_speech_metrics(answer: str, duration_seconds: int) -> dict:
    """Analyze speech delivery metrics from answer text and duration."""
    words = answer.split()
    word_count = len(words)
    duration_minutes = duration_seconds / 60.0
    
    # Speaking pace (words per minute)
    wpm = word_count / duration_minutes if duration_minutes > 0 else 0
    
    # Filler word detection
    filler_words = ["um", "uh", "like", "you know", "so", "well", "actually", "basically", "literally", "totally"]
    filler_count = sum(1 for word in words if word.lower().strip('.,!?') in filler_words)
    filler_rate = filler_count / word_count if word_count > 0 else 0
    
    # Sentence clarity (simple metrics)
    sentences = re.split(r'[.!?]+', answer)
    sentences = [s.strip() for s in sentences if s.strip()]
    sentence_count = len(sentences)
    avg_words_per_sentence = word_count / sentence_count if sentence_count > 0 else 0
    
    # Clarity score (arbitrary, higher is better)
    clarity_score = min(10, max(0, 10 - (avg_words_per_sentence - 15) * 0.5 - filler_rate * 20))
    
    return {
        "word_count": word_count,
        "duration_seconds": duration_seconds,
        "speaking_pace_wpm": round(wpm, 1),
        "filler_word_count": filler_count,
        "filler_rate": round(filler_rate, 3),
        "sentence_count": sentence_count,
        "avg_words_per_sentence": round(avg_words_per_sentence, 1),
        "sentence_clarity_score": round(clarity_score, 1)
    }

# ===========================================================================
# AUTH HELPERS
# ===========================================================================
def get_current_user(request: Request, db: Session):
    username = request.cookies.get("user")
    if not username:
        return None
    return db.query(User).filter(User.username == username).first()

def hash_password(password: str):
    # PBKDF2-SHA512 has no length limit; full password is always hashed
    return pwd_context.hash(password)


def _is_bcrypt_hash(hashed: str) -> bool:
    return isinstance(hashed, str) and hashed.startswith("$2")


def verify_password(plain: str, hashed: str):
    try:
        if _is_bcrypt_hash(hashed):
            # bcrypt only supports passwords up to 72 bytes.
            plain = plain[:72]
        # Verify against hashed password (works with both PBKDF2 and legacy bcrypt)
        return pwd_context.verify(plain, hashed)
    except Exception as e:
        logger.warning(f"Password verification error: {e}")
        return False


def get_or_create_user_profile(db: Session, user: User) -> UserProfile:
    """
    Ensure a UserProfile row exists for the given auth user.

    This keeps creation logic in one place and can be reused from
    login, resume upload, and interview flows.
    """
    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
    if profile:
        return profile

    profile = UserProfile(
        user_id=user.id,
        email=user.username,  # treat username as email by default
    )
    db.add(profile)
    db.commit()
    db.refresh(profile)
    return profile


# ---------------------------------------------------------------------------
# User Skill Profile helpers (DB-backed skill vector)
# ---------------------------------------------------------------------------

def get_skill_profile(db: Session, user_id: int):
    """Fetch the user skill profile row, or None if it doesn't exist."""
    return db.query(UserSkillProfileRow).filter(UserSkillProfileRow.user_id == user_id).first()


def create_skill_profile(db: Session, user_id: int) -> UserSkillProfileRow:
    """Create a new skill profile row for a user (no-op if already exists)."""
    profile = get_skill_profile(db, user_id)
    if profile:
        return profile

    profile = UserSkillProfileRow(
        user_id=user_id,
        technical_skills={},
        interview_skills={},
        communication_skills={},
        overall_score=0.0,
        interview_count=0,
    )
    db.add(profile)
    db.commit()
    db.refresh(profile)
    return profile
def update_skill_profile(db: Session, user_id: int, skill_data: dict):
    profile = get_skill_profile(db, user_id)

    if not profile:
        profile = create_skill_profile(db, user_id)

    if "technical_skills" in skill_data:
        profile.technical_skills = skill_data["technical_skills"]

    if "interview_skills" in skill_data:
        profile.interview_skills = skill_data["interview_skills"]

    if "communication_skills" in skill_data:
        profile.communication_skills = skill_data["communication_skills"]

    if "overall_score" in skill_data:
        profile.overall_score = float(skill_data["overall_score"])

    if "interview_count" in skill_data:
        profile.interview_count = int(skill_data["interview_count"])

    profile.last_updated = datetime.utcnow()

    db.add(profile)
    db.commit()
    db.refresh(profile)

    return profile
# ===========================================================================
# PAGE ROUTES
# ===========================================================================
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse(request, "home.html", {"request": request})

@app.get("/signup", response_class=HTMLResponse)
def signup_page(request: Request):
    return templates.TemplateResponse(request, "signup.html", {"request": request})

@app.post("/signup")
def signup(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    if db.query(User).filter(User.username == username).first():
        return templates.TemplateResponse(request, "signup.html", {"request": request, "message": "User already exists"})
    hashed_password = hash_password(password)
    user = User(username=username, password=hashed_password)
    db.add(user)
    db.commit()
    return RedirectResponse("/login", status_code=303)

@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse(request, "login.html", {"request": request})

@app.post("/login")
def login(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.username == username).first()

    if not user:
        return templates.TemplateResponse(request, "login.html", {"request": request, "message": "Invalid credentials"}
        )

    if not verify_password(password, user.password):
        return templates.TemplateResponse(request, "login.html", {"request": request, "message": "Invalid credentials"}
        )

    # Ensure a profile row exists for this user.
    get_or_create_user_profile(db, user)

    resp = RedirectResponse("/index", status_code=303)
    resp.set_cookie(key="user", value=username, httponly=True)
    return resp

@app.get("/logout")
def logout(request: Request, db: Session = Depends(get_db)):
    resp = RedirectResponse("/")
    resp.delete_cookie("user")
    return resp

@app.post("/update-resume")
def update_resume(request: Request, file: UploadFile, db: Session = Depends(get_db)):

    user = get_current_user(request, db)

    if not user:
        return RedirectResponse("/login", status_code=303)

    os.makedirs(UPLOAD_DIR, exist_ok=True)

    file_location = os.path.join(UPLOAD_DIR, f"{user.id}_{file.filename}")

    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    if profile:
        profile.resume_file_path = file_location
        db.commit()

    return RedirectResponse("/profile", status_code=303)

@app.get("/index", response_class=HTMLResponse)
def index(request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()

    role = profile.role_applied_for if profile else ""
    level = profile.current_designation if profile else ""
    company_name = profile.company_name if profile else ""
    job_description = profile.job_description if profile else ""

    return templates.TemplateResponse(request, "index.html", {
            "request": request,
            "username": user.username,
            "saved_role": role,
            "saved_level": level,
            "saved_company_name": company_name or "",
            "saved_job_description": job_description or "",
        },
    )

@app.get("/progress", response_class=HTMLResponse)
@app.get("/progress/")
def progress_redirect(request: Request):
    """Legacy progress page removed — send users to the profile dashboard."""
    return RedirectResponse("/profile", status_code=303)

# ===========================================================================
# RESUME UPLOAD
# ===========================================================================
@app.post("/api/save_interview_inputs")
async def save_interview_inputs(
    request: Request,
    db: Session = Depends(get_db),
):
    """Save role/level/company/JD without requiring a resume upload."""
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    body = await request.json()
    role = str(body.get("role", "")).strip()
    level = str(body.get("level", "")).strip()
    company_name = str(body.get("company_name", "")).strip()
    job_description = str(body.get("job_description", "")).strip()

    if not role or not level:
        raise HTTPException(status_code=400, detail="role and level are required")

    profile = _save_interview_profile_fields(
        db,
        user,
        role=role,
        level=level,
        company_name=company_name,
        job_description=job_description,
    )
    return {
        "status": "ok",
        "role": profile.role_applied_for,
        "level": profile.current_designation,
        "company_name": profile.company_name or "",
        "job_description_length": len(profile.job_description or ""),
    }


@app.post("/api/upload_resume")
async def upload_resume(
    request: Request,
    file: UploadFile = File(...),
    role: str = Form(...),
    level: str = Form(...),
    company_name: str = Form(""),
    job_description: str = Form(""),
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    content = await file.read()
    original_name = file.filename or "resume"
    filename = original_name.lower()
    text = ""

    if filename.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif filename.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)
    else:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Could not extract text from the file")

    # Persist raw text in memory for question generation
    resume_store[user.username] = text

    # Persist file to disk for later download from profile
    upload_dir = Path("uploads")
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_name = f"user_{user.id}_{int(time.time())}_{original_name}"
    file_path = upload_dir / safe_name
    with open(file_path, "wb") as f:
        f.write(content)

    # Create / update structured skill profile using the resume
    profile_row = get_or_create_user_profile(db, user)

    # Ask LLM for compact skill profile using the new prompt manager
    try:
        parsed = await generate_resume_skill_profile(text, role, level)
    except Exception:
        parsed = {}

    skills = parsed.get("skills") if isinstance(parsed, dict) else None
    if not isinstance(skills, list):
        skills = []
    skills = [str(s).strip() for s in skills if str(s).strip()]

    skill_gaps = parsed.get("skill_gaps") if isinstance(parsed, dict) else None
    if not isinstance(skill_gaps, list):
        skill_gaps = []
    skill_gaps = [str(s).strip() for s in skill_gaps if str(s).strip()]

    improvement_suggestions = (
        parsed.get("improvement_suggestions") if isinstance(parsed, dict) else ""
    ) or ""

    strength_pct = 0.0
    if isinstance(parsed, dict):
        try:
            strength_pct = float(parsed.get("strength_percentage", 0.0))
        except (TypeError, ValueError):
            strength_pct = 0.0

    # Build Pydantic UserSkillProfile
    basic_info = BasicUserInfo(
        user_id=str(user.id),
        name=user.username,
        target_role=role,
        experience_level=level,
        resume_file_path=str(file_path),
    )
    resume_data = ResumeData(skills=skills)
    profile_obj: UserSkillProfile = create_user_profile(basic_info, resume_data)

    # Detect weaknesses based on initial profile (may be empty at this stage)
    detect_weaknesses(profile_obj)

    profile_row.role_applied_for = role
    profile_row.current_designation = level
    profile_row.company_name = company_name.strip() or None
    profile_row.job_description = job_description.strip() or None
    profile_row.resume_file_path = str(file_path)
    profile_row.extracted_skills = json.dumps(skills)
    profile_row.skill_strength_percentage = strength_pct
    profile_row.skill_gaps = json.dumps(skill_gaps)
    profile_row.improvement_suggestions = improvement_suggestions
    profile_row.profile_json = profile_obj.json()
    profile_row.updated_at = datetime.utcnow()

    db.add(profile_row)
    db.commit()
    return {
        "status": "ok",
        "length": len(text),
        "preview": text[:200],
        "skills": skills,
        "skill_gaps": skill_gaps,
        "strength_percentage": strength_pct,
    }

# ===========================================================================
# AUDIO TRANSCRIPTION (Whisper)
# ===========================================================================
@app.post("/api/transcribe")
async def transcribe_endpoint(file: UploadFile = File(...)):
    audio_bytes = await file.read()

    # Save as webm (actual format from browser)
    temp_input = f"temp_audio_{datetime.now().timestamp()}.webm"
    temp_output = temp_input.replace(".webm", ".wav")

    with open(temp_input, "wb") as f:
        f.write(audio_bytes)

    try:
        # Convert webm → wav using your existing ffmpeg function
        _convert_audio(temp_input, temp_output)

        # Transcribe the converted wav
        text = transcribe_audio(temp_output)

    finally:
        # Clean up temp files
        if os.path.exists(temp_input):
            os.remove(temp_input)
        if os.path.exists(temp_output):
            os.remove(temp_output)

    return {"transcript": text}
# ===========================================================================
# INTERVIEW — START (returns page for voice interview)
# ===========================================================================
def _build_interview_context(
    request: Request,
    user,
    role: str,
    level: str,
    course_id: int | None,
    db: Session,
    company_name: str = "",
    job_description: str = "",
    interview_mode: str = "placement_simulation",
    selected_round: str | None = None,
):
    completed_modules = []
    course_topics = []

    if course_id is not None:
        course = db.query(Course).filter(Course.id == course_id).first()
        if course and course.user_id == user.id:
            modules = (
                db.query(Module)
                .filter(Module.course_id == course_id)
                .order_by(Module.order_index.asc())
                .all()
            )
            completed_modules = [m.title for m in modules if m.is_completed]
            course_topics = [m.title for m in modules if m.title]

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
    effective_company = (company_name or (profile.company_name if profile else "") or "").strip()
    effective_jd = (job_description or (profile.job_description if profile else "") or "").strip()

    if company_name or job_description:
        _save_interview_profile_fields(
            db,
            user,
            role=role,
            level=level,
            company_name=effective_company,
            job_description=effective_jd,
        )

    interview_sessions[user.username] = {
        "role": role,
        "level": level,
        "course_id": course_id,
        "company_name": effective_company,
        "job_description": effective_jd,
        "questions": [],
        "answers": [],
        "completed_modules": completed_modules,
        "course_topics": course_topics,
        "categories": [],
        **_init_interview_session_fields(interview_mode, selected_round),
    }
    session = interview_sessions[user.username]
    _prepare_interview_session(
        session,
        interview_mode=interview_mode,
        selected_round=selected_round,
    )

    current_round = session.get("current_round", "technical")
    round_executable = current_round in EXECUTABLE_ROUNDS

    logger.info(
        "Interview entry: mode=%s round=%s executable=%s",
        session.get("interview_mode"),
        current_round,
        round_executable,
    )

    return templates.TemplateResponse(request, "interview.html", {
            "request": request,
            "username": user.username,
            "user_id": user.id,
            "role": role,
            "level": level,
            "course_id": course_id,
            "interview_mode": session.get("interview_mode"),
            "current_round": current_round,
            "round_executable": round_executable,
        },
    )

@app.post("/start_interview/", response_class=HTMLResponse)
def start_interview_page(
    request: Request,
    role: str = Form(...),
    level: str = Form(...),
    company_name: str = Form(""),
    job_description: str = Form(""),
    course_id: int | None = Form(None),
    interview_mode: str = Form("placement_simulation"),
    selected_round: str = Form("technical"),
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    return _build_interview_context(
        request,
        user,
        role,
        level,
        course_id,
        db,
        company_name=company_name,
        job_description=job_description,
        interview_mode=interview_mode,
        selected_round=selected_round,
    )

@app.get("/interview/start", response_class=HTMLResponse)
def start_interview_get(
    request: Request,
    role: str | None = None,
    level: str | None = None,
    course_id: int | None = None,
    interview_mode: str | None = None,
    selected_round: str | None = None,
    db: Session = Depends(get_db),
):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    effective_role = (role or "").strip()
    effective_level = (level or "").strip()
    completed_modules: list[str] = []
    course_topics: list[str] = []

    if course_id is not None:
        course = db.query(Course).filter(Course.id == course_id).first()
        if course and course.user_id == user.id:
            effective_role = (course.role or effective_role or "").strip()
            effective_level = (course.level or effective_level or "").strip()
            modules = (
                db.query(Module)
                .filter(Module.course_id == course_id)
                .order_by(Module.order_index.asc())
                .all()
            )
            completed_modules = [m.title for m in modules if m.is_completed and m.title]
            course_topics = [m.title for m in modules if m.title][:8]

    profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
    if not effective_role and profile and profile.role_applied_for:
        effective_role = profile.role_applied_for
    if not effective_level and profile and profile.current_designation:
        effective_level = profile.current_designation

    if not effective_role:
        effective_role = "Software Engineer"
    if not effective_level:
        effective_level = "Junior"

    resume_text = (resume_store.get(user.username) or "").strip()
    if not resume_text and profile and profile.resume_file_path:
        resume_path = Path(profile.resume_file_path)
        if resume_path.exists():
            try:
                suffix = resume_path.suffix.lower()
                if suffix == ".pdf":
                    from pypdf import PdfReader
                    reader = PdfReader(str(resume_path))
                    resume_text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
                elif suffix == ".docx":
                    from docx import Document
                    doc = Document(str(resume_path))
                    resume_text = "\n".join(p.text for p in doc.paragraphs).strip()
                else:
                    resume_text = resume_path.read_text(encoding="utf-8", errors="ignore").strip()
            except Exception:
                resume_text = ""
    if resume_text:
        resume_store[user.username] = resume_text

    effective_company = (profile.company_name if profile else "") or ""
    effective_jd = (profile.job_description if profile else "") or ""

    interview_sessions[user.username] = {
        "role": effective_role,
        "level": effective_level,
        "course_id": course_id,
        "company_name": effective_company,
        "job_description": effective_jd,
        "questions": [],
        "answers": [],
        "completed_modules": completed_modules,
        "course_topics": course_topics,
        "categories": [],
        **_init_interview_session_fields(
            interview_mode or "placement_simulation",
            selected_round,
        ),
    }
    session = interview_sessions[user.username]
    _prepare_interview_session(
        session,
        interview_mode=interview_mode or "placement_simulation",
        selected_round=selected_round,
    )

    current_round = session.get("current_round", "technical")
    round_executable = current_round in EXECUTABLE_ROUNDS

    logger.info(
        "Interview entry: mode=%s round=%s executable=%s",
        session.get("interview_mode"),
        current_round,
        round_executable,
    )

    return templates.TemplateResponse(request, "interview.html", {
            "request": request,
            "username": user.username,
            "user_id": user.id,
            "role": effective_role,
            "level": effective_level,
            "course_id": course_id,
            "interview_mode": session.get("interview_mode"),
            "current_round": current_round,
            "round_executable": round_executable,
        },
    )


# ===========================================================================
# INTERVIEW API — Retrieve next interview question using conversation history
# ===========================================================================
@app.post("/api/interview/next_question")
async def api_interview_next_question(request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    body = await request.json()

    user_id = body.get("user_id")
    history = body.get("history", [])

    # -----------------------------
    # VALIDATION
    # -----------------------------
    if user_id is None:
        raise HTTPException(status_code=400, detail="Missing user_id")

    if str(user_id) != str(user.id) and str(user_id) != str(user.username):
        raise HTTPException(status_code=403, detail="Invalid user_id")

    if not isinstance(history, list):
        raise HTTPException(status_code=400, detail="history must be a list")

    # -----------------------------
    # EXTRACT MEMORY FROM HISTORY
    # -----------------------------
    previous_questions = [
        h.get("question", "").strip()
        for h in history
        if isinstance(h, dict) and h.get("question")
    ]

    # -----------------------------
    # SESSION SETUP
    # -----------------------------
    session = interview_sessions.get(user.username)

    if session is None:
        interview_sessions[user.username] = {
            "role": body.get("role", "Software Engineer"),
            "level": body.get("level", "Junior"),
            "questions": [],
            "answers": [],
            "categories": [],
            **_init_interview_session_fields("placement_simulation"),
        }
        session = interview_sessions[user.username]
        _prepare_interview_session(session, interview_mode="placement_simulation")

    _ensure_round_fields(session)

    # Ensure categories always exist
    if "categories" not in session:
        session["categories"] = []

    used_categories = session.get("categories", [])

    # -----------------------------
    # CONTEXT
    # -----------------------------
    role = session.get("role", body.get("role", "Software Engineer"))
    level = session.get("level", body.get("level", "Junior"))
    resume_text = resume_store.get(user.username, "")
    completed_modules = session.get("completed_modules", [])
    course_topics = session.get("course_topics", [])

    if "company_name" not in session or "job_description" not in session:
        profile = db.query(UserProfile).filter(UserProfile.user_id == user.id).first()
        session.setdefault("company_name", (profile.company_name if profile else "") or "")
        session.setdefault("job_description", (profile.job_description if profile else "") or "")

    await _ensure_session_jd_analysis(session, resume_text)
    jd_fields = _interview_jd_payload(session)

    # -----------------------------
    # LLM CALL (FIXED INPUT)
    # -----------------------------
    course_id = session.get("course_id", body.get("course_id"))
    llm_payload = {
        "role": role,
        "level": level,
        "resume_text": resume_text,
        "completed_modules": completed_modules,
        "course_topics": course_topics,
        "previous_questions": previous_questions,
        "used_categories": used_categories,
        **jd_fields,
    }
    if course_id is not None:
        llm_payload["context_priority"] = "course_focused"
        llm_payload["instruction_override"] = "Generate mostly questions from course_topics and completed_modules. Only 20-30% can be resume or general role-based."
        logger.info("Course-focused interview context applied for user_id=%s course_id=%s", user.id, course_id)

    result = await question_chain.invoke(llm_payload)

    if result.status != "success":
        raise HTTPException(status_code=500, detail="Failed to generate next interview question")

    # -----------------------------
    # PARSE RESPONSE SAFELY
    # -----------------------------
    try:
        parsed = extract_json(result.output)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse LLM response: {str(e)}")

    if not isinstance(parsed, dict) or "question" not in parsed:
        raise HTTPException(status_code=500, detail="Invalid question response format")

    # -----------------------------
    # UPDATE SESSION STATE
    # -----------------------------
    if "category" in parsed and parsed["category"]:
        session["categories"].append(parsed["category"])

    session["questions"].append(parsed.get("question", ""))

    # -----------------------------
    # RETURN RESPONSE
    # -----------------------------
    return JSONResponse(content=parsed)

# ===========================================================================
# INTERVIEW API — Batch evaluate all answers after interview ends
# ===========================================================================
@app.post("/api/interview/evaluate")
async def api_interview_evaluate(request: Request, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")


    body = await request.json()
    
    if "questions_answers" in body:
        # ---------------------------------------------------------------
        # Full analysis pipeline (Features 1-8)
        # ---------------------------------------------------------------
        questions_answers = body.get("questions_answers", [])
        if not questions_answers:
            raise HTTPException(status_code=400, detail="No questions_answers provided")


        session = interview_sessions.get(user.username, {})
        _ensure_round_fields(session)
        role = body.get("role", session.get("role", "Software Developer"))
        level = body.get("level", session.get("level", "mid"))
        company_name = session.get("company_name", "")
        job_description = session.get("job_description", "")
        n = len(questions_answers)

        # -- FEATURE 2: Speech Delivery Analysis -----------------------
        speech_analyses: list[dict] = []
        for qa in questions_answers:
            # Compute duration from timestamps when available
            duration = float(qa.get("duration", 60))
            start_time = qa.get("start_time", "")
            end_time = qa.get("end_time", "")
            if start_time and end_time:
                try:
                    st = datetime.fromisoformat(start_time.replace("Z", "+00:00"))
                    et = datetime.fromisoformat(end_time.replace("Z", "+00:00"))
                    duration = max((et - st).total_seconds(), 1.0)
                except Exception:
                    pass
            speech_analyses.append(
                analyze_speech_delivery(qa.get("answer", ""), duration)
            )

        # Aggregate voice metrics
        avg_pace = sum(a["speaking_pace_wpm"] for a in speech_analyses) / max(n, 1)
        total_fillers = sum(a["filler_word_count"] for a in speech_analyses)
        avg_clarity = sum(a["clarity_score"] for a in speech_analyses) / max(n, 1)
        avg_engagement = sum(a["engagement_score"] for a in speech_analyses) / max(n, 1)

        # -- FEATURE 4: Confidence Score -------------------------------
        confidence = compute_confidence_score(speech_analyses)

        # -- FEATURE 3: Content Analysis (LLM) ------------------------
        content_result = await evaluate_content(
            role,
            level,
            questions_answers,
            company_name=company_name,
            job_description=job_description,
        )
        content_answers = content_result.get("answers", [])
        aggregate = content_result.get("aggregate", {})

        content_scores = [a.get("score", 50) for a in content_answers]
        content_avg = sum(content_scores) / max(len(content_scores), 1)

        # -- FEATURE 6: Overall Score & Verdict ------------------------
        overall = compute_overall_score(content_avg, avg_clarity, avg_engagement,questions_answers)
        verdict = compute_recruiter_verdict(overall, role)

        # -- FEATURE 4 (cont.): Session Metadata -----------------------
        total_duration = sum(a["duration_seconds"] for a in speech_analyses)

        # -- FEATURE 1: Candidate Profile ------------------------------
        candidate_profile = {
            "role": role,
            "level": level,
            "interview_date": datetime.now(timezone.utc).astimezone(timezone(timedelta(hours=5, minutes=30))).strftime("%Y-%m-%d %H:%M IST"),
            "total_questions": n,
        }
        # -- FEATURE 7: Detailed Answer Report -------------------------

        detailed_answers: list[dict] = []

        for i, qa in enumerate(questions_answers):

            ca = content_answers[i] if i < len(content_answers) else {}
            sa = speech_analyses[i] if i < len(speech_analyses) else {}

            # Ensure question text is from input
            question_text = qa.get("question", "")
            answer_text = qa.get("answer", "")


            # Normalize strengths
            strengths = ca.get("strengths", [])

            if isinstance(strengths, str):
                strengths = [strengths]

            if not strengths:
                strengths = ["Answer attempted."]

            # Normalize weaknesses
            weaknesses = ca.get("weaknesses", [])

            if isinstance(weaknesses, str):
                weaknesses = [weaknesses]

            if not weaknesses:
                weaknesses = ["Needs improvement."]

            detailed_answers.append({
                "question": question_text,
                "transcript": answer_text,
                "score": ca.get("score", 50),
                "strengths": strengths,
                "weaknesses": weaknesses,
                "ideal_answer": ca.get("ideal_answer", "No ideal answer generated."),
                "weak_topics": ca.get("weak_topics", []),
               "voice_metrics": {
    "speaking_pace_wpm": sa.get("speaking_pace_wpm", 0),
    "filler_word_count": sa.get("filler_word_count", 0),
    "clarity_score": sa.get("clarity_score", 0),
    "engagement_score": sa.get("engagement_score", 0),
    "word_count": sa.get("word_count", 0),
}
            })

        # -- FEATURE 8: Assemble Report --------------------------------
        jd_analysis = session.get("jd_analysis")
        if not jd_analysis and job_description:
            resume_text = resume_store.get(user.username, "")
            jd_analysis = await analyze_resume_vs_jd(
                resume_text,
                job_description,
                llm_service=llm_service,
                prompt_manager=prompt_manager,
            )
            session["jd_analysis"] = jd_analysis

        report = {
            "candidate_profile": candidate_profile,
            "overall_score": overall,
            "verdict": verdict,
            "performance_summary": "",  # filled below by LLM
            "job_match_analysis": jd_analysis if job_description else None,
            "voice_analysis": {
                "speaking_pace_wpm": round(avg_pace, 1),
                "total_filler_words": total_fillers,
                "clarity_score": round(avg_clarity),
                "engagement_score": round(avg_engagement),
                "confidence_score": confidence,
                "total_duration_seconds": round(total_duration, 1),
                "per_answer": speech_analyses,
            },
            "content_analysis": {
                "average_score": round(content_avg),
                "relevance_score": aggregate.get("relevance_score", round(content_avg)),
                "depth_score": aggregate.get("depth_score", round(content_avg)),
                "star_method_score": aggregate.get("star_method_score", round(content_avg)),
                "structured_thinking_score": aggregate.get(
                    "structured_thinking_score", round(content_avg)
                ),
                "problem_solving_score": aggregate.get(
                    "problem_solving_score", round(content_avg)
                ),
            },
            "detailed_answers": detailed_answers,
        }

        # -- FEATURE 5: Performance Summary ----------------------------
        # Use overall_feedback from content evaluation to avoid a second LLM call.
        # Fall back to a separate LLM call only when the evaluator didn't provide one.
        overall_feedback = content_result.get("overall_feedback", "")
        if overall_feedback and overall_feedback.strip():
            report["performance_summary"] = overall_feedback.strip()
        else:
            summary = await generate_performance_summary(report)
            report["performance_summary"] = summary

        # Expose weak topics at the top level for report.html
        report["weak_topics"] = categorize_weak_topics(content_result.get("weak_topics", []))

        # ============================================================
        # RL: BANDIT-BASED COURSE RECOMMENDATION & REWARD FEEDBACK
        # ============================================================
        try:
            weak_topics = content_result.get("weak_topics", [])
            if isinstance(weak_topics, str):
                weak_topics = [weak_topics]
            weak_topics = [str(t).strip() for t in weak_topics if str(t).strip()]

            # STEP 1: Compute state
            state_id = get_state_id(overall, len(weak_topics))

            # STEP 2: Get user state for previous score and session tracking
            user_state = db.query(UserState).filter(UserState.user_id == user.id).first()
            previous_score = float(user_state.avg_score) if user_state and user_state.avg_score is not None else overall
            prev_state_id = user_state.state_id if user_state else state_id
            session_count = user_state.session_count if user_state else 0

            history_metrics = _get_recent_interview_metrics(user.id, db)
            last_action = history_metrics["last_action"]
            action_streak = history_metrics["action_streak"]
            previous_weak_topics = history_metrics["previous_weak_topics"]
            previous_confidence = history_metrics["previous_confidence"]
            persistent_weak_topics = history_metrics["persistent_weak_topics"]

            log_bandit_state(
                state_id=state_id,
                overall_score=overall,
                weak_topics=weak_topics,
                weak_count=len(weak_topics),
                previous_score=previous_score,
                prev_state_id=prev_state_id,
                session_count=session_count,
            )

            # STEP 3: Select action using bandit (terminal logs inside select_action)
            course_learner = ContextualBandit(db=db, action_space="course")
            action = course_learner.select_action(
                state_id,
                user_state,
                last_action=last_action,
                consecutive_action_count=action_streak,
            )

            # STEP 4: Map action → difficulty (simple lookup)
            course_difficulty = _ACTION_TO_DIFFICULTY.get(action, "medium")

            # STEP 5: Create course with real learning resources
            new_course_id = None
            fallback_used = False
            try:
                new_course_id = await create_course_from_resources(
                    user,
                    role,
                    level,
                    weak_topics,
                    action,
                    db,
                    difficulty=course_difficulty,
                    company=company_name,
                    interview_score=overall,
                )
            except Exception as course_error:
                logger.warning(
                    f"Course creation failed for action={action}: {str(course_error)}"
                )
                new_course_id = None

            # Fallback: retry with simpler params if first attempt failed
            if not new_course_id:
                try:
                    fallback_used = True
                    new_course_id = await create_course_from_resources(
                        user,
                        role,
                        level,
                        _prioritize_top_weak_topics(weak_topics, 2),
                        "revision",
                        db,
                        difficulty="easy",
                        company=company_name,
                        interview_score=overall,
                    )
                except Exception as fallback_error:
                    logger.warning(f"Fallback course creation failed: {str(fallback_error)}")
                    new_course_id = None

            # STEP 6: Fetch course details for response
            new_course = None
            if new_course_id:
                new_course_row = db.query(Course).filter(Course.id == new_course_id).first()
                if new_course_row:
                    new_course = {
                        "course_id": new_course_row.id,
                        "title": new_course_row.title,
                    }

            # STEP 7: Calculate adaptive reward for learning progression
            score_improvement = overall - previous_score  # range: -100 to +100
            reward = calculate_reward(
                current_score=overall,
                previous_score=previous_score,
                current_weak_topics=weak_topics,
                previous_weak_topics=previous_weak_topics,
                current_confidence=confidence,
                previous_confidence=previous_confidence,
            )

            weak_topic_progress = 0.0
            if previous_weak_topics:
                prev_set = set(t.lower().strip() for t in previous_weak_topics if t)
                curr_set = set(t.lower().strip() for t in weak_topics if t)
                overlap = len(prev_set.intersection(curr_set))
                weak_topic_progress = max(0.0, min(1.0, (len(prev_set) - overlap) / 3.0))

            confidence_improvement = 0.0
            if previous_confidence is not None:
                confidence_improvement = (max(0.0, min(100.0, confidence)) - max(0.0, min(100.0, previous_confidence))) / 100.0

            log_reward_calculation(
                overall_score=overall,
                previous_score=previous_score,
                score_improvement=score_improvement,
                reward=reward,
                weak_topic_progress=weak_topic_progress,
                confidence_improvement=confidence_improvement,
            )

            course_title = None
            if new_course_id:
                _course_row = db.query(Course).filter(Course.id == new_course_id).first()
                course_title = _course_row.title if _course_row else None

            log_course_generation_decision(
                action=action,
                course_topics=weak_topics,
                course_difficulty=course_difficulty,
                weak_topics=weak_topics,
                new_course_id=new_course_id,
                course_title=course_title,
                fallback_used=fallback_used,
            )

            # STEP 8: Update bandit with reward (terminal logs inside update_action_value)
            course_learner.update_action_value(prev_state_id, action, reward)

            # STEP 9: Update user state for next session
            if user_state:
                user_state.state_id = state_id
                user_state.last_score = overall
                user_state.avg_score = (
                    user_state.avg_score * user_state.session_count + overall
                ) / (user_state.session_count + 1)
                user_state.current_proficiency = user_state.avg_score / 100.0
                user_state.weak_topics = weak_topics
                user_state.session_count += 1
            else:
                user_state = UserState(
                    user_id=user.id,
                    state_id=state_id,
                    weak_topics=weak_topics,
                    avg_score=overall,
                    last_score=overall,
                    current_proficiency=overall / 100.0,
                    session_count=1,
                )
                db.add(user_state)

            db.flush()

            # STEP 10: Populate report with course and action
            report["new_course_id"] = new_course_id
            if new_course is not None:
                report["new_course"] = new_course
            report["recommended_action"] = action
            report["rl_metrics"] = {
                "state": state_id,
                "previous_state": prev_state_id,
                "action": action,
                "course_topics": weak_topics,
                "course_difficulty": course_difficulty,
                "reward": reward,
                "score_improvement": score_improvement,
                "new_course_id": new_course_id,
            }

            log_bandit_complete(
                state_id=state_id,
                action=action,
                reward=reward,
                new_course_id=new_course_id,
            )
        except Exception as e:
            logger.warning(f"RL course recommendation failed: {str(e)}", exc_info=True)
            report["new_course_id"] = None
            report["recommended_action"] = "revision"

        # Store report for the /report page (latest only)
        report_store[user.username] = report
        interview_row = Interview(
            user_id=user.id,
            role=role,
            date=datetime.utcnow(),
            score=overall,
            report_json=json.dumps(report),
        )
        db.add(interview_row)
        db.commit()
        db.refresh(interview_row)

        # Update rich UserSkillProfile based on this interview
        profile_row = get_or_create_user_profile(db, user)
        profile_obj: UserSkillProfile
        try:
            if profile_row.profile_json:
                profile_obj = UserSkillProfile.model_validate_json(profile_row.profile_json)
            else:
                # Minimal fallback profile if none existed yet
                basic_info = BasicUserInfo(
                    user_id=str(user.id),
                    name=user.username,
                    target_role=role,
                    experience_level=level,
                    resume_file_path=profile_row.resume_file_path or "",
                )
                resume_data = ResumeData(skills=[])
                profile_obj = create_user_profile(basic_info, resume_data)
        except Exception:
            basic_info = BasicUserInfo(
                user_id=str(user.id),
                name=user.username,
                target_role=role,
                experience_level=level,
                resume_file_path=profile_row.resume_file_path or "",
            )
            resume_data = ResumeData(skills=[])
            profile_obj = create_user_profile(basic_info, resume_data)

        # Record interview per question into profile
        for i, qa in enumerate(questions_answers):
            ca = content_answers[i] if i < len(content_answers) else {}
            topic = role  # Treat role as the main topic/skill for now

            sb = ScoreBreakdown(
                correctness=float(ca.get("score", 50)),
                conceptual_depth=float(ca.get("score", 50)),
                clarity=float(avg_clarity),
                feedback="",
            )
            rec = InterviewRecord(
                question=qa.get("question", ""),
                topic=topic,
                answer_transcript=qa.get("answer", ""),
                evaluation_score=float(ca.get("score", 50)),
                score_breakdown=sb,
            )
            record_interview_result(profile_obj, rec)

        # Sync weak topics from LLM result - handle both dict and list
        weak_from_llm = []
        weak_topics_raw = content_result.get("weak_topics", [])
        if isinstance(weak_topics_raw, dict):
            for topics in weak_topics_raw.values():
                if isinstance(topics, list):
                    weak_from_llm.extend(topics)
                else:
                    weak_from_llm.append(str(topics) if topics else "")
        elif isinstance(weak_topics_raw, list):
            weak_from_llm.extend(weak_topics_raw)
        
        # Safe access to weak_topics with fallback
        profile_weak_topics = getattr(profile_obj, 'weak_topics', [])
        if not isinstance(profile_weak_topics, list):
            profile_weak_topics = []
        
        for w in weak_from_llm:
            if not w:
                continue
            w_str = str(w).strip() if w else ""
            if not w_str:
                continue
            w_norm = w_str.lower()
            if w_norm not in {t.lower() for t in profile_weak_topics if t}:
                profile_weak_topics.append(w_str)
        
        # Update profile with clean weak_topics list
        profile_obj.weak_topics = profile_weak_topics

        # Recompute weaknesses & recommend micro-courses
        detect_weaknesses(profile_obj)

        # --------- Update skill vector based on interview metrics ---------
        def _int(v, default=0):
            try:
                return int(round(float(v)))
            except Exception:
                return default

        voice = report.get("voice_analysis", {})
        content = report.get("content_analysis", {})

        def _metric(v, fallback=0):
            try:
                return float(v)
            except (TypeError, ValueError):
                return float(fallback)

        interview_metrics = {
            "relevance": _metric(content.get("relevance_score"), content_avg),
            "explanation_depth": _metric(content.get("depth_score"), content_avg),
            "problem_solving": _metric(
                content.get("problem_solving_score"), content_avg
            ),
            "structured_thinking": _metric(
                content.get("structured_thinking_score"), content_avg
            ),
            "star_method": _metric(content.get("star_method_score"), content_avg),
            "clarity": _metric(voice.get("clarity_score", 0)),
            "confidence": _metric(voice.get("confidence_score", 0)),
            "engagement": _metric(voice.get("engagement_score", 0)),
            "speaking_pace": _pace_to_skill_score(_metric(voice.get("speaking_pace_wpm", 0))),
            "filler_control": max(
                0.0,
                100.0 - min(100.0, _metric(voice.get("total_filler_words", 0)) * 3),
            ),
        }

        # Update the stored skill vector using the interview metrics
        skill_vector = UserSkillVector(
            technical_skills=profile_obj.technical_skills,
            interview_skills=profile_obj.interview_skills,
            communication_skills=profile_obj.communication_skills,
        )
        skill_vector = update_skill_vector(skill_vector, interview_metrics)

        profile_obj.technical_skills = skill_vector.technical_skills
        profile_obj.interview_skills = skill_vector.interview_skills
        profile_obj.communication_skills = skill_vector.communication_skills
        profile_obj.overall_score = calculate_overall_score(skill_vector)

        # Aggregate simple fields for quick display in profile page
        skills_list = []
        avg_score = overall

        if hasattr(profile_obj, "skill_graph") and profile_obj.skill_graph:
            try:
                skills_list = [node.skill_name for node in profile_obj.skill_graph.values() if hasattr(node, 'skill_name')]
                if profile_obj.skill_graph:
                    scores = [n.score for n in profile_obj.skill_graph.values() if hasattr(n, 'score')]
                    if scores:
                        avg_score = sum(scores) / len(scores)
            except Exception:
                pass

        profile_row.role_applied_for = role
        profile_row.current_designation = level
        profile_row.extracted_skills = json.dumps(skills_list)
        profile_row.skill_strength_percentage = float(profile_obj.overall_score)
        profile_row.skill_gaps = json.dumps(getattr(profile_obj, 'weak_topics', []))
        # Use performance summary as latest improvement suggestions snapshot
        profile_row.improvement_suggestions = report.get("performance_summary", "")
        profile_row.profile_json = profile_obj.json()
        profile_row.updated_at = datetime.utcnow()

        # Persist a compact skill vector record for quick access / querying.
        update_skill_profile(
            db,
            user.id,
            {
                "technical_skills": profile_obj.technical_skills.dict(),
                "interview_skills": profile_obj.interview_skills.dict(),
                "communication_skills": profile_obj.communication_skills.dict(),
                "overall_score": profile_obj.overall_score,
                "interview_count": profile_obj.interview_count,
            },
        )

        # Return the updated skill profile as part of the report payload.
        report["skill_profile"] = profile_obj.dict()

        db.add(profile_row)

        # Persist answers to DB
        for i, qa in enumerate(questions_answers):
            attempt = InterviewAttempt(
                user_id=user.id,
                role=role,
                topic="voice-interview",
                difficulty=level,
                answer=qa.get("answer", ""),
                feedback="",
            )
            db.add(attempt)

        # Update skill progress
        skill = (
            db.query(SkillProgress)
            .filter(SkillProgress.user_id == user.id, SkillProgress.skill == role)
            .first()
        )
        if not skill:
            skill = SkillProgress(
                user_id=user.id, skill=role, attempts=1, weak=overall < 50
            )
            db.add(skill)
        else:
            skill.attempts += 1
            skill.weak = overall < 50

        db.commit()

        return report
    
    else:
        # Old format
        role = body.get("role", "")
        answers = body.get("answers", [])  # [{question, answer}, ...]

        if not answers:
            raise HTTPException(status_code=400, detail="No answers provided")

        # Store answers in session
        session = interview_sessions.get(user.username, {})
        session["answers"] = answers

        evaluation = await evaluate_answers(role, answers)

        # Normalize top-level keys (AI may use variant names)
        if isinstance(evaluation, dict):
            for alt in ("evaluations", "feedback", "responses", "results"):
                if alt in evaluation and "feedback_per_question" not in evaluation:
                    evaluation["feedback_per_question"] = evaluation.pop(alt)
                    break
            for alt in ("tips", "suggestions", "general_tips"):
                if alt in evaluation and "improvement_tips" not in evaluation:
                    evaluation["improvement_tips"] = evaluation.pop(alt)
                    break
            for alt in ("resources", "recommended_resources", "study_resources"):
                if alt in evaluation and "learning_resources" not in evaluation:
                    evaluation["learning_resources"] = evaluation.pop(alt)
                    break

        # Validate and build fallback if parsing failed or structure is wrong
        if not isinstance(evaluation, dict) or "feedback_per_question" not in evaluation:
            evaluation = {
                "feedback_per_question": [
                    {
                        "question": qa.get("question", ""),
                        "candidate_answer": qa.get("answer", ""),
                        "feedback": "",
                        "improved_answer": "",
                    }
                    for qa in answers
                ],
                "improvement_tips": [
                    "Practice structuring your answers with concrete examples.",
                    "Use the STAR method (Situation, Task, Action, Result) for behavioral questions.",
                    "Research the company and role thoroughly before interviews.",
                ],
                "learning_resources": [{"topic": "Interview Preparation", "resource": "Practice common behavioral and technical questions for your role."}],
            }

        # Normalize per-question field names (AI may use variants)
        for i, fb in enumerate(evaluation.get("feedback_per_question", [])):
            qa = answers[i] if i < len(answers) else {}
            for src, dst in [
                ("answer", "candidate_answer"),
                ("user_answer", "candidate_answer"),
                ("better_answer", "improved_answer"),
                ("suggested_answer", "improved_answer"),
                ("sample_answer", "improved_answer"),
                ("ideal_answer", "improved_answer"),
            ]:
                if src in fb and dst not in fb:
                    fb[dst] = fb.pop(src)
            fb.setdefault("question", qa.get("question", ""))
            fb.setdefault("candidate_answer", qa.get("answer", ""))
            fb.setdefault("feedback", "")
            fb.setdefault("improved_answer", "")

        # Normalize learning_resources to always be list of {topic, resource}
        raw_resources = evaluation.get("learning_resources", [])
        normalized = []
        for r in raw_resources:
            if isinstance(r, dict):
                topic = r.get("topic") or r.get("name") or r.get("title") or ""
                resource = r.get("resource") or r.get("description") or r.get("link") or r.get("url") or ""
                normalized.append({"topic": topic, "resource": resource})
            elif isinstance(r, str):
                normalized.append({"topic": r, "resource": ""})
        evaluation["learning_resources"] = normalized

        # Persist each answer to DB
        for i, qa in enumerate(answers):
            attempt = InterviewAttempt(
                user_id=user.id,
                role=role,
                topic="voice-interview",
                difficulty="adaptive",
                answer=qa.get("answer", ""),
                feedback="",
            )
            db.add(attempt)

        # Update skill progress
        skill = db.query(SkillProgress).filter(SkillProgress.user_id == user.id, SkillProgress.skill == role).first()
        if not skill:
            skill = SkillProgress(user_id=user.id, skill=role, attempts=1, weak=False)
            db.add(skill)
        else:
            skill.attempts += 1

        db.commit()

        return evaluation

# ===========================================================================
# REPORT PAGE
# ===========================================================================
def _get_resources_by_concept(course_id: int, db: Session) -> dict:
    """Return CourseResource rows grouped by concept for a given course."""
    resources = (
        db.query(CourseResource)
        .filter(CourseResource.course_id == course_id)
        .order_by(CourseResource.rank_score.desc())
        .all()
    )
    grouped: dict = {}
    for cr in resources:
        grouped.setdefault(cr.concept or "General", []).append(cr)
    return grouped


@app.get("/report", response_class=HTMLResponse)
def report_page(request: Request, db: Session = Depends(get_db)):
    """Render the full interview performance report."""
    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")
    report = report_store.get(user.username)
    if not report:
        return RedirectResponse("/index")

    course_id = report.get("new_course_id")
    resources_by_concept = _get_resources_by_concept(course_id, db) if course_id else {}

    return templates.TemplateResponse(request, "report.html", {
            "request": request,
            "username": user.username,
            "report": report,
            "resources_by_concept": resources_by_concept,
        },
    )


@app.get("/interview-report/{interview_id}", response_class=HTMLResponse)
def interview_report_page(
    request: Request,
    interview_id: int,
    db: Session = Depends(get_db),
):
    """
    View a past interview report stored in DB (secure).
    """

    from models import Interview
    import json

    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")

    # ✅ IMPORTANT: restrict to user's own reports
    interview = (
        db.query(Interview)
        .filter(
            Interview.id == interview_id,
            Interview.user_id == user.id
        )
        .first()
    )

    if not interview:
        raise HTTPException(status_code=404, detail="Report not found")

    try:
        report = json.loads(interview.report_json)
    except Exception:
        report = {}

    course_id = report.get("new_course_id")
    resources_by_concept = _get_resources_by_concept(course_id, db) if course_id else {}

    return templates.TemplateResponse(request, "report.html", {
            "request": request,
            "username": user.username,
            "report": report,
            "resources_by_concept": resources_by_concept,
        },
    )


def _pace_to_skill_score(wpm: float) -> float:
    """Map words-per-minute to 0–100 skill scale."""
    if wpm <= 0:
        return 50.0
    if 120 <= wpm <= 165:
        return 92.0
    if 90 <= wpm < 120 or 165 < wpm <= 190:
        return 72.0
    return 55.0



_COURSE_ACTION_LABEL = {
    "revision": "Revision",
    "easy": "Foundation",
    "mixed": "Applied",
    "advanced": "Mastery",
}



def _make_adaptive_course_title(
    db: Session,
    user_id: int,
    role: str,
    action: str,
    difficulty: str,
    topics: list[str],
) -> str:
    """Deterministic, unique course title from RL decision + weak topics."""
    action_label = _COURSE_ACTION_LABEL.get(action, action.title())
    primary = (topics[0] if topics else role).strip()[:40]
    base = f"{difficulty.title()} {action_label}: {primary}"
    existing = {
        row[0]
        for row in db.query(Course.title).filter(Course.user_id == user_id).all()
        if row[0]
    }
    title = base
    n = 2
    while title in existing:
        title = f"{base} · Session {n}"
        n += 1
    return title


def _enrich_profile_from_interviews(profile_data: dict, interviews: list) -> dict:
    """Pull latest real scores from stored interview reports (avoid static 50s)."""
    if not interviews:
        profile_data["has_interview_data"] = False
        profile_data["interview_count"] = 0
        return profile_data

    profile_data["has_interview_data"] = True
    profile_data["interview_count"] = len(interviews)
    latest = max(interviews, key=lambda iv: iv.date or datetime.min)
    try:
        report = json.loads(latest.report_json) if latest.report_json else {}
    except Exception:
        report = {}

    sp = report.get("skill_profile") if isinstance(report.get("skill_profile"), dict) else {}
    for group in ("interview_skills", "communication_skills", "technical_skills"):
        if group in sp and isinstance(sp[group], dict):
            base = profile_data.get(group) if isinstance(profile_data.get(group), dict) else {}
            profile_data[group] = {**base, **sp[group]}

    if sp.get("overall_score") is not None:
        profile_data["overall_score"] = sp["overall_score"]
    elif report.get("overall_score") is not None:
        profile_data["overall_score"] = report["overall_score"]
    elif latest.score is not None:
        profile_data["overall_score"] = float(latest.score)

    return profile_data


def _compute_working_skill_gaps(profile_data: dict) -> list:
    """Gaps from interview + communication skills only (excludes unused technical DSA/OS/CN)."""
    gaps = []
    for skills in (
        profile_data.get("interview_skills") or {},
        profile_data.get("communication_skills") or {},
    ):
        if not isinstance(skills, dict):
            continue
        for key, val in skills.items():
            try:
                score = float(val)
            except (TypeError, ValueError):
                score = 0
            if score < 60:
                gaps.append({
                    "name": key.replace("_", " ").title(),
                    "score": round(score, 1),
                })
    gaps.sort(key=lambda x: x["score"])
    return gaps


def _extract_report_weak_topics(report: dict) -> list[str]:
    """Flatten weak topics from stored interview report JSON."""
    topics: list[str] = []
    weak = report.get("weak_topics")
    if isinstance(weak, dict):
        for value in weak.values():
            if isinstance(value, list):
                topics.extend(str(t).strip() for t in value if t)
            elif value:
                topics.append(str(value).strip())
    elif isinstance(weak, list):
        topics = [str(t).strip() for t in weak if t]

    seen = set()
    deduped = []
    for t in topics:
        key = t.lower()
        if key and key not in seen:
            seen.add(key)
            deduped.append(t)
    return deduped[:12]


def _course_summary_from_entry(course_entry: dict) -> dict:
    modules = course_entry.get("modules") or []
    total = len(modules) or 1
    completed = sum(1 for m in modules if m.get("is_completed"))
    return {
        "course_id": course_entry["course_id"],
        "title": course_entry.get("title") or "Adaptive Course",
        "level": course_entry.get("level"),
        "progress_pct": round((completed / total) * 100),
        "completed_modules": completed,
        "total_modules": total,
        "is_complete": completed == total and total > 0,
        "modules": modules,
        "created_at": course_entry.get("created_at"),
        "course_url": f"/course/{course_entry['course_id']}",
    }


def _build_course_history(db: Session, user_id: int) -> list[dict]:
    """Fetch all courses with module progress for the user."""
    course_history = []
    courses = (
        db.query(Course)
        .filter(Course.user_id == user_id)
        .order_by(Course.created_at.asc())
        .all()
    )
    for course in courses:
        modules = (
            db.query(Module)
            .filter(Module.course_id == course.id)
            .order_by(Module.order_index.asc())
            .all()
        )
        module_entries = []
        for module in modules:
            attempt = (
                db.query(ModuleAttempt)
                .filter(
                    ModuleAttempt.user_id == user_id,
                    ModuleAttempt.module_id == module.id,
                )
                .order_by(ModuleAttempt.created_at.desc())
                .first()
            )
            score = attempt.score if attempt else None
            attempts = 0
            is_passed = False
            total_questions = attempt.total_questions if attempt else None
            if attempt:
                payload = attempt.answers
                if isinstance(payload, dict):
                    try:
                        attempts = int(payload.get("attempt_count", 1))
                    except Exception:
                        attempts = 1
                    is_passed = bool(payload.get("is_passed", False))
                elif isinstance(payload, list):
                    attempts = 1
                    is_passed = bool(score is not None and score >= 2)

            module_entries.append(
                {
                    "module_id": module.id,
                    "title": module.title,
                    "order_index": module.order_index,
                    "is_unlocked": module.is_unlocked,
                    "is_completed": module.is_completed,
                    "score": score,
                    "total_questions": total_questions,
                    "is_passed": is_passed,
                    "attempts": attempts,
                }
            )

        course_history.append(
            {
                "course_id": course.id,
                "title": course.title,
                "role": course.role,
                "level": course.level,
                "status": course.status,
                "created_at": course.created_at.strftime("%Y-%m-%d") if course.created_at else None,
                "modules": module_entries,
            }
        )
    return course_history


def _build_unified_learning_timeline(
    interviews: list,
    course_history: list,
    profile_row,
) -> list[dict]:
    """
    Role → interviews → expandable adaptive course / follow-up details.
    Courses not linked to an interview appear as standalone entries under the role.
    """
    from urllib.parse import quote

    course_by_id = {c["course_id"]: c for c in course_history}
    linked_course_ids: set[int] = set()
    roles_map: dict[str, dict] = {}

    def get_role_bucket(role: str) -> dict:
        key = (role or "general").strip().lower() or "general"
        display = (role or "General").strip()
        if not display or display.lower() == "general":
            display = "General"
        else:
            display = display.title()
        if key not in roles_map:
            roles_map[key] = {
                "role_key": key,
                "role_display": display,
                "interviews": [],
            }
        return roles_map[key]

    sorted_interviews = sorted(
        interviews,
        key=lambda iv: iv.date or datetime.min,
        reverse=True,
    )

    for iv in sorted_interviews:
        role = (iv.role or "General").strip()
        bucket = get_role_bucket(role)
        report: dict = {}
        try:
            report = json.loads(iv.report_json) if iv.report_json else {}
        except Exception:
            report = {}

        rl = report.get("rl_metrics") or {}
        new_course_id = report.get("new_course_id") or rl.get("new_course_id")
        adaptive_course = None
        if new_course_id and new_course_id in course_by_id:
            adaptive_course = _course_summary_from_entry(course_by_id[new_course_id])
            linked_course_ids.add(int(new_course_id))

        weak_topics = _extract_report_weak_topics(report)
        candidate = report.get("candidate_profile") or {}
        level = candidate.get("level") or (
            profile_row.current_designation if profile_row else "Junior"
        )
        role_for_url = candidate.get("role") or role

        followup_url = None
        if adaptive_course and adaptive_course["is_complete"]:
            followup_url = (
                f"/interview/start?role={quote(str(role_for_url))}"
                f"&level={quote(str(level))}"
                f"&course_id={adaptive_course['course_id']}"
            )

        summary = report.get("performance_summary") or ""
        if len(summary) > 220:
            summary = summary[:217].rstrip() + "..."

        bucket["interviews"].append(
            {
                "id": iv.id,
                "date": iv.date.strftime("%Y-%m-%d") if iv.date else "",
                "date_display": iv.date.strftime("%b %d, %Y") if iv.date else "",
                "score": round(float(iv.score), 1) if iv.score is not None else 0,
                "weak_topics": weak_topics,
                "summary_snippet": summary,
                "rl_action": report.get("recommended_action") or rl.get("action"),
                "rl_difficulty": rl.get("course_difficulty"),
                "rl_topics": rl.get("course_topics") or [],
                "rl_reward": rl.get("reward"),
                "report_url": f"/interview-report/{iv.id}",
                "adaptive_course": adaptive_course,
                "followup_url": followup_url,
                "level": level,
                "role_for_url": role_for_url,
            }
        )

    standalone_courses: list[dict] = []
    for course in course_history:
        if course["course_id"] in linked_course_ids:
            continue
        standalone_courses.append(_course_summary_from_entry(course))

    timeline = sorted(
        roles_map.values(),
        key=lambda r: (r["interviews"][0]["date"] if r["interviews"] else ""),
        reverse=True,
    )
    return timeline, standalone_courses


@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        return RedirectResponse("/login")

    from models import Interview
    import json

    # Ensure profile exists
    profile_row = get_or_create_user_profile(db, user)

    profile_data = {
        "interview_skills": {},
        "communication_skills": {},
        "technical_skills": {},
        "overall_score": 0,
        "interview_count": 0,
    }
    if profile_row.profile_json:
        try:
            loaded = json.loads(profile_row.profile_json)
            if isinstance(loaded, dict):
                profile_data.update(loaded)
        except Exception:
            pass

    # -------------------------------
    # EXTRACTED SKILLS
    # -------------------------------
    extracted_skills = []
    if profile_row.extracted_skills:
        try:
            extracted_skills = json.loads(profile_row.extracted_skills)
        except:
            extracted_skills = []

    # -------------------------------
    # SKILL GAPS
    # -------------------------------
    skill_gaps = []
    if profile_row.skill_gaps:
        try:
            skill_gaps = json.loads(profile_row.skill_gaps)
        except:
            skill_gaps = []

    # -------------------------------
    # FETCH INTERVIEW HISTORY
    # -------------------------------
    interviews = (
        db.query(Interview)
        .filter(Interview.user_id == user.id)
        .order_by(Interview.date.asc())
        .all()
    )

    # -------------------------------
    # BUILD TIMELINE (for graphs)
    # -------------------------------
    timeline_by_role = {}

    for iv in interviews:
        role = iv.role.strip().lower()

        if role not in timeline_by_role:
            timeline_by_role[role] = []

        timeline_by_role[role].append({
            "date": iv.date.strftime("%Y-%m-%d"),
            "score": iv.score
        })

    # -------------------------------
    # IMPROVEMENT MESSAGE
    # -------------------------------
    improvement_message = (
        profile_row.improvement_suggestions
        if profile_row and profile_row.improvement_suggestions
        else "Keep practicing interviews to improve your profile."
    )

    profile_data = _enrich_profile_from_interviews(profile_data, interviews)

    course_history = _build_course_history(db, user.id)
    learning_timeline, standalone_courses = _build_unified_learning_timeline(
        interviews, course_history, profile_row
    )
    dashboard_analytics = build_dashboard_analytics(profile_data, interviews)

    return templates.TemplateResponse(request, "profile.html", {
            "request": request,
            "username": user.username,
            "user_profile": profile_row,
            "skill_profile": profile_data,
            "dashboard": dashboard_analytics,
            "timeline_by_role": timeline_by_role,
            "improvement_message": improvement_message,
            "learning_timeline": learning_timeline,
            "standalone_courses": standalone_courses,
        },
    )

@app.get("/interview-history", response_class=HTMLResponse)
def interview_history_redirect(request: Request):
    """
    Simple semantic route that redirects to the profile page where
    interview history is rendered.
    """
    return RedirectResponse("/profile")


# Deprecated: replaced by DB-based course system
# @app.get("/api/course/stream")
# async def api_course_stream(request: Request, role: str, level: str, db: Session = Depends(get_db)):
#     user = get_current_user(request, db)
#     if not user:
#         raise HTTPException(status_code=401, detail="Not logged in")
#
#     course = Course(
#         user_id=user.id,
#         role=role,
#         title=f"{level.title()} {role} Course",
#         description="",
#         level=level,
#         status="draft"
#     )
#     db.add(course)
#     db.flush()
#
#     async def event_generator():
#         import asyncio
#
#         # Heartbeat
#         yield {"event": "heartbeat", "data": "connected"}
#
#         # =========================
#         # STAGE 1: OUTLINE
#         # =========================
#         yield {"event": "status", "data": "Generating course outline..."}
#
#         await asyncio.sleep(0.5)
#
#         outline_prompt_input = {
#             "skill": role,
#             "level": level,
#             "duration_hours": 20
#         }
#
#         # ===== OUTLINE =====
#         outline_prompt = prompt_manager.get_prompt("course_outline", **outline_prompt_input)
#         outline_text = await llm_service.invoke(outline_prompt)
#
#         try:
#             outline_json = extract_json(outline_text)
#         except:
#             match = re.search(r"\{.*\}", outline_text, re.DOTALL)
#             if match:
#                 outline_json = json.loads(match.group(0))
#             else:
#                 raise ValueError("Invalid outline JSON")
#
#         course.title = outline_json.get("course_title", course.title)
#         course.description = outline_json.get("description", "")
#         db.commit()
#
#         # Normalize modules and persist skeletons
#         modules = []
#         module_records = []
#         for idx, mod in enumerate(outline_json.get("modules", [])):
#             title = mod.get("module_title") or mod.get("title")
#             description = ", ".join(mod.get("topics", [])) if mod.get("topics") else mod.get("description", "")
#
#             module_record = Module(
#                 course_id=course.id,
#                 title=title,
#                 description=description,
#                 order_index=idx,
#                 is_unlocked=(idx == 0)
#             )
#             db.add(module_record)
#             db.flush()
#             module_records.append(module_record)
#
#             modules.append({
#                 "id": module_record.id,
#                 "title": title,
#                 "description": description
#             })
#
#         db.commit()
#
#         # Send outline
#         yield {
#             "event": "outline",
#             "data": json.dumps({
#                 "course_title": outline_json.get("course_title"),
#                 "description": outline_json.get("description"),
#                 "modules": modules
#             })
#         }
#
#         # =========================
#         # STAGE 2: MODULES (Deep Dive)
#         # =========================
#         for i, mod in enumerate(modules):
#             if await request.is_disconnected():
#                 return
#
#             yield {
#                 "event": "status",
#                 "data": f"Generating module {i+1}/{len(modules)}..."
#             }
#
#             module_prompt_input = {
#                 "skill": role,
#                 "module": mod["title"],
#                 "level": level
#             }
#
#             # ===== MODULE =====
#             module_prompt = prompt_manager.get_prompt("course_module_detail", **module_prompt_input)
#             module_text = await llm_service.invoke(module_prompt)
#
#             try:
#                 module_json = json.loads(module_text)
#             except:
#                 match = re.search(r"\{.*\}", module_text, re.DOTALL)
#                 if match:
#                     module_json = json.loads(match.group(0))
#                 else:
#                     raise ValueError("Invalid module JSON")
#
#             # Safety check
#             if len(module_json.get("quiz", [])) != 3:
#                 module_json["quiz"] = []
#
#             module_record = module_records[i]
#             module_record.content = module_json.get("content_markdown", "")
#             module_record.quiz = module_json.get("quiz", [])
#             db.commit()
#
#             yield {
#                 "event": "module_detail",
#                 "data": json.dumps({
#                     "module_id": module_record.id,
#                     "index": i,
#                     "module_title": module_json.get("module_title", module_record.title),
#                     "content_markdown": module_record.content,
#                     "quiz": module_record.quiz,
#                     "external_practice_tasks": module_json.get("external_practice_tasks", [])
#                 })
#             }
#
#         # =========================
#         # DONE
#         # =========================
#         yield {"event": "done", "data": "Course generation complete!"}
#
#     return EventSourceResponse(event_generator())
# ===========================================================================
# SKELETON-FIRST COURSE GENERATION SYSTEM (NEW ARCHITECTURE)
# ===========================================================================

@app.post("/api/course/create")
async def create_course(
    request: Request,
    role: str = Form(...),
    level: str = Form(...),
    db: Session = Depends(get_db),
):
    """
    Create a course with real learning resources from the retrieval pipeline.
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    try:
        weak_topics = [role]  # Use role as default topic for manual creation
        course_id = await create_course_from_resources(
            user, role, level, weak_topics, "manual", db, difficulty="medium",
        )

        if not course_id:
            raise HTTPException(status_code=500, detail="Failed to create course")

        return RedirectResponse(url=f"/course/{course_id}", status_code=303)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Course creation failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to create course")


# ===========================================================================
# Action → difficulty mapping (replaces old _compose_course_topics)
# ===========================================================================
_ACTION_TO_DIFFICULTY = {
    "revision": "easy",
    "easy": "easy",
    "mixed": "medium",
    "advanced": "hard",
}


async def create_course_from_resources(
    user,
    role: str,
    level: str,
    weak_topics: list[str],
    action: str,
    db: Session,
    difficulty: str = "medium",
    company: str = "",
    interview_score: float = 0.0,
) -> int | None:
    """
    Create a course backed by real learning resources (retrieval pipeline).

    Instead of generating fake LLM content, this fetches, ranks, and persists
    real learning resources from YouTube, GeeksforGeeks, etc.

    Args:
        user: User object
        role: Job role
        level: Experience level
        weak_topics: Weak topics from interview evaluation
        action: Bandit action (revision, easy, mixed, advanced, manual)
        db: Database session
        difficulty: Difficulty level (easy, medium, hard)
        company: Target company name (improves query specificity)
        interview_score: Overall interview score 0-100 (used to derive resource depth)
    """
    weak_topics = [str(t).strip() for t in weak_topics if str(t).strip()]
    weak_topics = weak_topics or ["core concepts"]

    title = _make_adaptive_course_title(
        db, user.id, role, action, difficulty, weak_topics
    )
    topics_preview = ", ".join(weak_topics[:3])
    description = f"Learning resources for {role} — focused on: {topics_preview}"

    course = Course(
        user_id=user.id,
        role=role,
        title=title,
        description=description,
        level=difficulty,
        status="generated",
    )
    db.add(course)
    db.flush()

    # Retrieve and rank resources using the pipeline
    if resource_pipeline is not None:
        try:
            # Derive resource depth from interview score when available
            if interview_score >= 75:
                resource_difficulty = "advanced"
            elif interview_score >= 45:
                resource_difficulty = "intermediate"
            elif interview_score > 0:
                resource_difficulty = "beginner"
            else:
                resource_difficulty = difficulty  # fall back to bandit difficulty

            ranked_results = await resource_pipeline.retrieve_and_rank(
                weak_concepts=weak_topics[:5],
                role=role,
                top_k=5,
                level=level,
                company=company,
                difficulty=resource_difficulty,
            )

            # Generate LLM explanations for retrieved resources.
            # The LLM re-ranks and explains — it never invents resources.
            explanations: dict = {}
            try:
                explanations = await resource_pipeline.generate_explanations(
                    ranked_results=ranked_results,
                    role=role,
                    level=level,
                    difficulty=resource_difficulty,
                )
            except Exception as exp_err:
                logger.warning(f"Explanation generation skipped: {exp_err}")

            for concept, ranked_list in ranked_results.items():
                concept_explanations = explanations.get(concept, {})
                for resource, score in ranked_list:
                    cr = CourseResource(
                        course_id=course.id,
                        user_id=user.id,
                        concept=concept,
                        resource_id=resource.id,
                        title=resource.title,
                        url=resource.url,
                        source=resource.source,
                        rank_score=round(score, 4),
                        explanation=concept_explanations.get(resource.id) or None,
                        resource_metadata=resource.to_dict(),
                    )
                    db.add(cr)

            total = sum(len(v) for v in ranked_results.values())
            logger.info(
                f"Created course {course.id} with {total} resources "
                f"for user {user.id} (action={action})"
            )
        except Exception as e:
            logger.warning(f"Resource retrieval failed for course {course.id}: {str(e)}")
    else:
        logger.warning("Resource pipeline not initialized, course created without resources")

    db.commit()
    return course.id


@app.get("/course/{course_id}", response_class=HTMLResponse)
def course_page(request: Request, course_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    course = db.query(Course).filter(Course.id == course_id).first()
    if not course:
        raise HTTPException(status_code=404, detail="Course not found")
    if course.user_id != user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    # Fetch learning resources grouped by concept
    resources_by_concept = _get_resources_by_concept(course.id, db)

    # Fallback: check for old module-based course (backward compat)
    modules = []
    if not resources_by_concept:
        modules = (
            db.query(Module)
            .filter(Module.course_id == course.id)
            .order_by(Module.order_index.asc())
            .all()
        )

    return templates.TemplateResponse(request, "course.html", {
            "request": request,
            "username": user.username,
            "course": course,
            "resources_by_concept": resources_by_concept,
            "modules": modules,
        },
    )


@app.get("/module/{module_id}", response_class=HTMLResponse)
def module_page(request: Request, module_id: int, db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    module = db.query(Module).filter(Module.id == module_id).first()
    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    course = db.query(Course).filter(Course.id == module.course_id).first()
    if not course or course.user_id != user.id:
        raise HTTPException(status_code=403, detail="Not authorized")

    if not module.is_unlocked:
        raise HTTPException(status_code=403, detail="Module not unlocked yet")

    return templates.TemplateResponse(request, "module.html", {
            "request": request,
            "username": user.username,
            "module_id": module.id,
            "module_title": module.title,
            "course_id": course.id,
            "role": course.role,
            "level": course.level,
            "is_final": module.is_final,
        },
    )

@app.get("/api/module/{module_id}")
async def get_module(module_id: int, request: Request, db: Session = Depends(get_db)):

    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    module = db.query(Module).filter(Module.id == module_id).first()

    if not module:
        raise HTTPException(status_code=404, detail="Module not found")

    def _strip_practice_markers(content_text: str):
        """Remove any embedded module practice-links marker block from stored content.

        We no longer expose or render practice links separately; this helper
        simply strips the marker block if present so the visible content is clean.
        """
        marker_start = "<!-- MODULE_PRACTICE_LINKS_START"
        marker_end = "MODULE_PRACTICE_LINKS_END -->"
        if marker_start in content_text and marker_end in content_text:
            try:
                clean_content = content_text.split(marker_start, 1)[0].strip()
                return clean_content
            except Exception:
                return content_text
        return content_text

    # Check if module is unlocked
    if not module.is_unlocked:
        raise HTTPException(status_code=403, detail="Module locked")

    # ✅ RETURN CACHED CONTENT (strip any historical practice-link markers)
    if module.content:
        logger.info("Using cached module")
        content_text = _strip_practice_markers(module.content)
        return {
            "module_id": module.id,
            "module_title": module.title,
            "content_markdown": content_text,
            "quiz": module.quiz if isinstance(module.quiz, list) else [],
        }

    course = db.query(Course).filter(Course.id == module.course_id).first()


    prompt = prompt_manager.get_prompt(
        "course_module_detail",
        skill=course.role,
        module=module.title,
        level=course.level,
        is_final=module.is_final,
        previous_modules=""
    )

    raw = await llm_service.invoke(prompt)

    try:
        module_json = extract_json(raw)
    except Exception as e:
        # Log parsing issue and attempt a tolerant fallback instead of failing hard.
        logger.warning(f"Module parse error for {module_id}: {str(e)}")
        logger.debug(f"Raw LLM response (for debugging): {repr(raw)}")

        # Fallback: try a loose regex-based extraction and a relaxed json.loads
        module_json = {}
        try:
            m = re.search(r"\{.*\}", raw, re.DOTALL)
            if m:
                candidate = m.group(0)
                # sanitize trailing commas
                candidate = re.sub(r',\s*}', '}', candidate)
                candidate = re.sub(r',\s*]', ']', candidate)
                module_json = json.loads(candidate)
        except Exception as e2:
            logger.warning(f"Fallback JSON parse also failed for module {module_id}: {e2}")
            module_json = {}

    raw_content = module_json.get("content_markdown", "") or "Detailed module content could not be retrieved. Please refresh the module."
    raw_content = _strip_practice_markers(raw_content)
    module.content = raw_content
    quiz_data = module_json.get("quiz", [])
    if isinstance(quiz_data, list) and len(quiz_data) == 3:
        valid = True
        for q in quiz_data:
            if not isinstance(q, dict):
                valid = False
                break
            options = q.get("options", [])
            answer = (q.get("answer") or "").upper()
            if not isinstance(options, list) or len(options) != 4 or answer not in ["A", "B", "C", "D"]:
                valid = False
                break
        if not valid:
            quiz_data = []
    else:
        quiz_data = []
    module.quiz = quiz_data

    # We no longer persist or return practice/external links for modules.
    # Ensure stored content remains clean and return only content + quiz.
    module.content = raw_content
    db.commit()
    logger.info("Module generated (links removed from flow)")

    return {
        "module_id": module.id,
        "module_title": module.title,
        "content_markdown": raw_content,
        "quiz": module.quiz,
    }

@app.post("/api/module/{module_id}/submit")
async def submit_quiz(
    module_id: int,
    request: Request,
    payload: dict = Body(...),
    db: Session = Depends(get_db)
):
    """
    Submit quiz answers for a module.
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")

    try:
        # Fetch module
        module = db.query(Module).filter(Module.id == module_id).first()
        if not module:
            raise HTTPException(status_code=404, detail="Module not found")

        # Check authorization
        course = db.query(Course).filter(Course.id == module.course_id).first()
        if course.user_id != user.id:
            raise HTTPException(status_code=403, detail="Not authorized")

        # Ensure quiz is loaded
        if not module.quiz or not isinstance(module.quiz, list) or len(module.quiz) != 3:
            raise HTTPException(status_code=400, detail="Module has no valid quiz")

        user_answers = payload.get("answers", [])
        if not isinstance(user_answers, list) or len(user_answers) != 3:
            raise HTTPException(status_code=400, detail="Must provide exactly 3 answers")

        # Score the quiz
        score = 0
        for user_ans, quiz_item in zip(user_answers, module.quiz):
            correct_ans = (quiz_item.get("answer") or "").upper()
            if isinstance(user_ans, str) and user_ans.upper() == correct_ans:
                score += 1

        passed = score >= 2  # At least 2 out of 3 correct
        # Store or update attempt
        attempt = (
            db.query(ModuleAttempt)
            .filter(
                ModuleAttempt.user_id == user.id,
                ModuleAttempt.module_id == module_id
            )
            .order_by(ModuleAttempt.created_at.desc())
            .first()
        )
        attempt_count = 1
        if attempt:
            existing_payload = attempt.answers if isinstance(attempt.answers, dict) else {}
            try:
                previous_count = int(existing_payload.get("attempt_count", 1))
            except Exception:
                previous_count = 1
            attempt_count = previous_count + 1
            attempt.score = score
            attempt.total_questions = 3
            attempt.answers = {
                "answers": user_answers,
                "is_passed": passed,
                "attempt_count": attempt_count,
            }
        else:
            attempt = ModuleAttempt(
                user_id=user.id,
                module_id=module_id,
                score=score,
                total_questions=3,
                answers={
                    "answers": user_answers,
                    "is_passed": passed,
                    "attempt_count": attempt_count,
                }
            )
            db.add(attempt)
        logger.info(
            "Module attempt persisted user_id=%s module_id=%s score=%s passed=%s attempts=%s",
            user.id,
            module_id,
            score,
            passed,
            attempt_count,
        )

        next_module_id = None
        # Mark module as completed and unlock next module if passed
        if passed:
            module.is_completed = True
            next_module = db.query(Module).filter(
                Module.course_id == module.course_id,
                Module.order_index == module.order_index + 1
            ).first()
            if next_module:
                next_module.is_unlocked = True
                next_module_id = next_module.id
                logger.info("Next module unlocked user_id=%s module_id=%s", user.id, next_module_id)

        db.commit()

        logger.info("Quiz submitted")

        return {
            "score": score,
            "passed": passed,
            "is_final": module.is_final,
            "next_module_id": next_module_id,
            "role": course.role,
            "level": course.level,
            "course_id": course.id
        }

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Quiz submission failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Quiz submission failed: {str(e)}")


@app.get("/api/course/{course_id}/status")
async def get_course_status(
    course_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    Get course completion status.
    
    - Check if all modules completed
    - Return: modules status, interview_unlocked flag
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    
    try:
        # Fetch course
        course = db.query(Course).filter(Course.id == course_id).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        
        # Check authorization
        if course.user_id != user.id:
            raise HTTPException(status_code=403, detail="Not authorized")
        
        # Get all modules
        modules = db.query(Module).filter(Module.course_id == course_id).order_by(Module.order_index).all()
        
        module_status = []
        all_completed = True
        for mod in modules:
            # Get best attempt score for this module
            best_attempt = db.query(ModuleAttempt).filter(
                ModuleAttempt.user_id == user.id,
                ModuleAttempt.module_id == mod.id
            ).order_by(ModuleAttempt.score.desc()).first()
            
            module_status.append({
                "id": mod.id,
                "title": mod.title,
                "order_index": mod.order_index,
                "is_unlocked": mod.is_unlocked,
                "is_completed": mod.is_completed,
                "best_score": best_attempt.score if best_attempt else None,
                "best_percentage": int((best_attempt.score / best_attempt.total_questions) * 100) if best_attempt else None
            })
            
            if not mod.is_completed:
                all_completed = False
        
        return {
            "course_id": course_id,
            "title": course.title,
            "role": course.role,
            "level": course.level,
            "modules": module_status,
            "all_completed": all_completed,
            "interview_unlocked": all_completed  # Unlock interview after all modules
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Course status check failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Course status check failed: {str(e)}")


async def generate_final_interview(
    user_id: int,
    course_id: int,
    db: Session
) -> dict:
    """
    Generate a final interview session after course completion.
    
    Uses:
    - Course role and level
    - User's resume
    - Weak topics (modules with low scores)
    
    Returns:
    - interview_session_id
    - initial questions
    """
    try:
        # Fetch course and user info
        course = db.query(Course).filter(Course.id == course_id).first()
        if not course:
            raise ValueError("Course not found")
        
        user_profile = db.query(UserProfile).filter(UserProfile.user_id == user_id).first()
        
        # Identify weak topics (modules with score < 70%)
        modules = db.query(Module).filter(Module.course_id == course_id).all()
        weak_topics = []
        
        for mod in modules:
            best_attempt = db.query(ModuleAttempt).filter(
                ModuleAttempt.user_id == user_id,
                ModuleAttempt.module_id == mod.id
            ).order_by(ModuleAttempt.score.desc()).first()
            
            if best_attempt:
                percentage = (best_attempt.score / best_attempt.total_questions) * 100
                if percentage < 70:
                    weak_topics.append(mod.title)
        
        # Create interview session
        session = InterviewSession(
            user_id=user_id,
            role=course.role,
            level=course.level,
            status="active"
        )
        db.add(session)
        db.flush()
        
        # Generate initial questions using interview chain
        resume_text = user_profile.resume_file_path if user_profile else ""
        
        initial_questions = await question_chain.invoke({
            "role": course.role,
            "level": course.level,
            "count": 5,
            "resume_text": resume_text,
            "previous_questions": [],
            "used_categories": []
        })
        
        db.commit()
        
        return {
            "interview_session_id": session.id,
            "role": course.role,
            "level": course.level,
            "weak_topics": weak_topics,
            "status": "started"
        }
    
    except Exception as e:
        db.rollback()
        logger.error(f"Final interview generation failed: {str(e)}")
        raise


@app.post("/api/course/{course_id}/start-interview")
async def start_final_interview(
    course_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """
    Trigger final interview generation after course completion.
    
    - Check all modules completed
    - Generate interview session
    - Return session ID
    """
    user = get_current_user(request, db)
    if not user:
        raise HTTPException(status_code=401, detail="Not logged in")
    
    try:
        # Verify course exists and belongs to user
        course = db.query(Course).filter(
            Course.id == course_id,
            Course.user_id == user.id
        ).first()
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        
        # Check if all modules completed
        modules = db.query(Module).filter(Module.course_id == course_id).all()
        all_completed = all(mod.is_completed for mod in modules)
        
        if not all_completed:
            raise HTTPException(status_code=400, detail="Not all modules completed")
        
        # Generate final interview
        interview_data = await generate_final_interview(user.id, course_id, db)
        
        return interview_data
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Interview start failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Interview start failed: {str(e)}")


# ===========================================================================
# RAG MONITORING ENDPOINTS
# ===========================================================================

@app.get("/api/rag/stats")
async def get_rag_stats():
    """Get RAG pipeline statistics and performance metrics."""
    
    if rag_pipeline is None:
        return JSONResponse(
            {"error": "RAG pipeline not initialized"},
            status_code=503
        )
    
    try:
        stats = rag_pipeline.get_stats()
        return {
            "status": "ok",
            "initialized": stats["initialized"],
            "documents_indexed": stats["vector_store"]["total_documents"],
            "retrievals_performed": stats["retrieval_stats"]["retrievals_performed"],
            "cache_hit_rate": (
                stats["retrieval_stats"]["cache_hits"] /
                (stats["retrieval_stats"]["cache_hits"] + stats["retrieval_stats"]["cache_misses"])
                if (stats["retrieval_stats"]["cache_hits"] + stats["retrieval_stats"]["cache_misses"]) > 0
                else 0
            ),
            "avg_retrieval_time_ms": stats["retrieval_stats"]["avg_retrieval_time"] * 1000,
            "cache_size": stats["cache_size"],
            "embedding_dimension": stats["vector_store"]["embedding_dimension"],
            "document_categories": stats["ingestion_stats"]["categories"]
        }
    
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500
        )


@app.post("/api/rag/clear-cache")
async def clear_rag_cache():
    """Clear the RAG retrieval cache."""
    
    if rag_pipeline is None:
        return JSONResponse(
            {"error": "RAG pipeline not initialized"},
            status_code=503
        )
    
    try:
        rag_pipeline.clear_cache()
        return {"status": "ok", "message": "RAG cache cleared"}
    
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500
        )


@app.get("/api/rag/test-retrieval")
async def test_rag_retrieval(query: str = "Python junior level interview"):
    """Test RAG retrieval with a sample query (for debugging)."""
    
    if rag_pipeline is None or not rag_pipeline.initialized:
        return JSONResponse(
            {"error": "RAG pipeline not initialized"},
            status_code=503
        )
    
    try:
        from services.rag.rag_config import RetrievalContext
        
        context = await rag_pipeline.retrieve_context(query)
        
        return {
            "status": "ok",
            "query": query,
            "retrieved_context": context[:500] + "..." if len(context) > 500 else context
        }
    
    except Exception as e:
        return JSONResponse(
            {"error": str(e)},
            status_code=500
        )


